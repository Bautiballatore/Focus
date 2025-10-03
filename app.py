from flask import Flask, render_template, request, redirect, session, url_for, flash, make_response, send_from_directory
# from flask_sqlalchemy import SQLAlchemy  # Comentado: no se usa en Supabase
from werkzeug.security import generate_password_hash, check_password_hash
from openai import OpenAI
from dotenv import load_dotenv
from supabase import create_client, Client
import os
import PyPDF2
import docx
import re
import time
import json
from datetime import datetime, timedelta
from io import BytesIO
import requests
import base64
import xml.etree.ElementTree as ET
import traceback

load_dotenv()

# Configuración de la aplicación
app = Flask(__name__, template_folder='Templates')
app.config.from_object('config.ProductionConfig' if os.environ.get('FLASK_ENV') == 'production' else 'config.DevelopmentConfig')
app.jinja_env.globals.update(range=range)

# Configuración de clave secreta para sesiones
app.secret_key = os.environ.get('SECRET_KEY', 'dev-secret-key-change-in-production')

# Configuración de sesiones seguras
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=7)  # Sesiones expiran en 7 días
app.config['SESSION_COOKIE_SECURE'] = False  # Permitir HTTP para desarrollo y Heroku
app.config['SESSION_COOKIE_HTTPONLY'] = True  # Previene acceso JavaScript a cookies
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'  # Protección CSRF básica para desarrollo
app.config['SESSION_COOKIE_DOMAIN'] = None  # No restringir dominio

# Context processor para pasar información del usuario a todos los templates
@app.context_processor
def inject_user():
    """Inyectar información del usuario en todos los templates"""
    # Verificar directamente la sesión de Flask
    has_user_id = 'user_id' in session
    has_user_email = 'user_email' in session
    is_auth = has_user_id and has_user_email
    
    print(f"🔍 Context processor - has_user_id: {has_user_id}, has_user_email: {has_user_email}, is_auth: {is_auth}")
    
    # Agregar fecha de hoy para los inputs de fecha
    today = datetime.now().strftime('%Y-%m-%d')
    
    if is_auth:
        # Obtener información adicional del usuario desde Supabase
        try:
            if supabase:
                user_response = supabase.table('usuarios').select('*').eq('id', session.get('user_id')).execute()
                if user_response.data:
                    user_data = user_response.data[0]
                    
                    # Obtener carpetas del usuario
                    carpetas_response = supabase.table('carpetas').select('id, nombre, color').eq('usuario_id', session.get('user_id')).order('nombre').execute()
                    carpetas = carpetas_response.data if carpetas_response.data else []
                    
                    return {
                        'current_user': {
                            'id': session.get('user_id'),
                            'email': session.get('user_email'),
                            'nombre': session.get('user_nombre'),
                            'fecha_registro': datetime.fromisoformat(user_data.get('fecha_registro', datetime.utcnow().isoformat()).replace('Z', '+00:00')),
                            'como_nos_conociste': user_data.get('como_nos_conociste'),
                            'plataforma_uso': user_data.get('plataforma_uso'),
                            'preguntas_completadas': user_data.get('preguntas_completadas', 0),
                            'total_examenes_rendidos': user_data.get('total_examenes_rendidos', 0),
                            'correctas_total': user_data.get('correctas_total', 0),
                            'parciales_total': user_data.get('parciales_total', 0),
                            'incorrectas_total': user_data.get('incorrectas_total', 0),
                            'ultima_actividad': user_data.get('ultima_actividad'),
                            'is_authenticated': True
                        },
                        'user_carpetas': carpetas,
                        'today': today
                    }
        except Exception as e:
            print(f"Error obteniendo datos del usuario para template: {e}")
        
        # Fallback si no se pueden obtener los datos adicionales
        print(f"🔄 Context processor - Fallback para usuario: {session.get('user_email')}")
        return {
            'current_user': {
                'id': session.get('user_id'),
                'email': session.get('user_email'),
                'nombre': session.get('user_nombre'),
                'fecha_registro': datetime.utcnow(),
                'como_nos_conociste': None,
                'plataforma_uso': None,
                'preguntas_completadas': 0,
                'total_examenes_rendidos': 0,
                'correctas_total': 0,
                'parciales_total': 0,
                'incorrectas_total': 0,
                'ultima_actividad': None,
                'is_authenticated': True
            },
            'user_carpetas': [],
            'today': today
        }
    
    print(f"❌ Context processor - Usuario no autenticado")
    return {
        'current_user': {
            'is_authenticated': False
        },
        'user_carpetas': [],
        'today': today
    }

# Inicializar Supabase
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_ANON_KEY = os.getenv("SUPABASE_ANON_KEY")

if SUPABASE_URL and SUPABASE_ANON_KEY:
    supabase: Client = create_client(SUPABASE_URL, SUPABASE_ANON_KEY)
    print("✅ Supabase client initialized successfully!")
else:
    print("❌ Supabase credentials not found!")
    supabase = None

# Configuración de Google OAuth
GOOGLE_CLIENT_ID = os.getenv("GOOGLE_CLIENT_ID")
GOOGLE_CLIENT_SECRET = os.getenv("GOOGLE_CLIENT_SECRET")

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

app_id = "AV6EGRRK9V"

@app.route("/")
def index():
    """Página principal que también maneja el callback de OAuth"""
    # Verificar si hay un código de autorización en la URL (callback de OAuth)
    code = request.args.get('code')
    
    if code:
        print(f"🔍 Código de autorización recibido en index: {code}")
        print(f"🔄 Procesando callback de Google OAuth desde index...")
        
        try:
            # Intentar obtener el usuario autenticado directamente
            print(f"🔍 Intentando obtener usuario con supabase.auth.get_user()...")
            current_user = supabase.auth.get_user()
            print(f"🔍 Tipo de current_user: {type(current_user)}")
            print(f"🔍 Valor de current_user: {current_user}")
            
            if current_user is None:
                print(f"❌ supabase.auth.get_user() devolvió None - OAuth no completado correctamente")
                print(f"🔍 Esto puede indicar un problema con la configuración de OAuth en Supabase")
                raise Exception("OAuth no completado - usuario no disponible")
            
            if hasattr(current_user, 'user') and current_user.user:
                print(f"✅ Usuario obtenido después de OAuth: {current_user.user.email}")
            else:
                print(f"❌ Usuario no tiene estructura válida")
                raise Exception("Estructura de usuario inválida")
            
            # Usar directamente la información de auth.users - mucho más simple
            print(f"✅ Usando información directa de auth.users")
            
            # Guardar usuario en sesión de Flask
            session['user_id'] = current_user.user.id
            session['user_email'] = current_user.user.email
            
            # Obtener nombre del user_metadata o usar email como fallback
            if hasattr(current_user.user, 'user_metadata') and isinstance(current_user.user.user_metadata, dict):
                session['user_nombre'] = current_user.user.user_metadata.get('nombre', current_user.user.email.split('@')[0])
            else:
                session['user_nombre'] = current_user.user.email.split('@')[0]
            
            print(f"✅ Usuario autenticado exitosamente: {current_user.user.email}")
            print(f"🔍 ID: {current_user.user.id}")
            print(f"🔍 Nombre: {session['user_nombre']}")
            
            # Para usuarios nuevos de Google, siempre ir a preguntas primero
            print(f"🔄 Usuario nuevo de Google - redirigiendo a preguntas de usuario")
            return redirect(url_for("preguntas_usuario"))
            
        except Exception as e:
            print(f"❌ Error obteniendo usuario después de OAuth: {e}")
            print(f"🔍 Tipo de error: {type(e)}")
            # Si falla, mostrar la página principal normalmente
            print(f"⚠️ Fallback: mostrando página principal sin autenticación")
    
    # Si no hay código o falló la autenticación, mostrar página principal normal
    return render_template("index.html")

@app.route("/favicon.ico")
def favicon():
    return redirect(url_for('static', filename='favicon.ico'))

# =====================================================
# NUEVAS RUTAS DE AUTENTICACIÓN CON SUPABASE AUTH
# =====================================================

# Lista de emails permitidos para registrarse
ALLOWED_EMAILS = [
    "admin@focusstudio.com",
    "test@focusstudio.com", 
    "demo@focusstudio.com",
    "usuario1@ejemplo.com",
    "usuario2@ejemplo.com",
    "estudiante@universidad.edu",
    "profesor@universidad.edu",
    # Agrega aquí los emails que quieras permitir
    # "nuevo@email.com",
]

def is_email_allowed(email):
    """
    Verifica si un email está en la lista de emails permitidos
    """
    return email.lower() in [e.lower() for e in ALLOWED_EMAILS]

def add_allowed_email(email):
    """
    Agrega un email a la lista de emails permitidos
    """
    if email not in ALLOWED_EMAILS:
        ALLOWED_EMAILS.append(email)
        print(f"✅ Email agregado a la lista: {email}")
        return True
    else:
        print(f"ℹ️ Email ya está en la lista: {email}")
        return False

def remove_allowed_email(email):
    """
    Remueve un email de la lista de emails permitidos
    """
    if email in ALLOWED_EMAILS:
        ALLOWED_EMAILS.remove(email)
        print(f"✅ Email removido de la lista: {email}")
        return True
    else:
        print(f"ℹ️ Email no está en la lista: {email}")
        return False

@app.route("/auth/signup", methods=["GET", "POST"])
def signup():
    if request.method == "POST":
        email = request.form["email"].lower()
        password = request.form["password"]
        nombre = request.form["nombre"]
        
        # Verificar si el email está en la lista de emails permitidos
        if not is_email_allowed(email):
            print(f"❌ Email {email} no está en la lista de emails permitidos")
            flash("Lo sentimos, tu email no está autorizado para registrarse en esta plataforma. Contacta al administrador si crees que esto es un error.")
            return render_template("registro.html")
        
        print(f"✅ Email {email} está permitido, procediendo con el registro...")
        
        try:
            if supabase:
                # Crear usuario con Supabase Auth
                user = supabase.auth.sign_up({
                    "email": email,
                    "password": password,
                    "options": {
                        "data": {
                            "nombre": nombre,
                            "preguntas_completadas": 0
                        }
                    }
                })
                
                if user.user:
                    print(f"✅ Usuario registrado exitosamente: {user.user.email}")
                    flash("Usuario registrado exitosamente. ")
                    return redirect(url_for('login'))
                else:
                    print(f"❌ No se pudo crear usuario")
                    flash("Error al registrar usuario. Intenta de nuevo.")
                    
        except Exception as e:
            print(f"❌ Error en signup: {e}")
            if "already registered" in str(e).lower():
                flash("El email ya está registrado. Por favor, usa otro email.")
            else:
                flash("Error al registrar usuario. Intenta de nuevo.")
                
        return render_template("registro.html")
    
    return render_template("registro.html")

@app.route("/admin/emails", methods=["GET", "POST"])
def manage_emails():
    """
    Ruta de administración para gestionar emails permitidos
    Solo para administradores
    """
    # Verificar si el usuario es administrador (puedes personalizar esta lógica)
    if not is_authenticated():
        return redirect(url_for('login'))
    
    # Verificar si es administrador (ejemplo: email específico)
    current_user_email = session.get('user_email', '')
    if current_user_email not in ['admin@focusstudio.com']:
        flash("No tienes permisos para acceder a esta página.")
        return redirect(url_for('index'))
    
    if request.method == "POST":
        action = request.form.get('action')
        email = request.form.get('email', '').strip()
        
        if action == 'add' and email:
            if add_allowed_email(email):
                flash(f"Email {email} agregado exitosamente.")
            else:
                flash(f"El email {email} ya está en la lista.")
        elif action == 'remove' and email:
            if remove_allowed_email(email):
                flash(f"Email {email} removido exitosamente.")
            else:
                flash(f"El email {email} no está en la lista.")
    
    return render_template("admin_emails.html", allowed_emails=ALLOWED_EMAILS)

@app.route("/auth/signin", methods=["GET", "POST"])
def signin():
    if request.method == "POST":
        email = request.form["email"].lower()
        password = request.form["password"]
        
        try:
            if supabase:
                # Iniciar sesión con Supabase Auth
                user = supabase.auth.sign_in_with_password({
                    "email": email,
                    "password": password
                })
                
                if user.user:
                    print(f"✅ Usuario autenticado exitosamente: {user.user.email}")
                    
                    # Verificar si el usuario existe en public.usuarios, si no, crearlo
                    try:
                        usuario_check = supabase.table('usuarios').select('id').eq('id', user.user.id).execute()
                        if not usuario_check.data:
                            print(f"🔄 Usuario no existe en public.usuarios, creándolo...")
                            nuevo_usuario = {
                                'id': user.user.id,
                                'email': user.user.email,
                                'nombre': user.user.user_metadata.get('nombre', email.split('@')[0]),
                                'fecha_registro': datetime.utcnow().isoformat(),
                                'preguntas_completadas': 0,
                                'total_examenes_rendidos': 0,
                                'correctas_total': 0,
                                'parciales_total': 0,
                                'incorrectas_total': 0,
                                'ultima_actividad': datetime.utcnow().isoformat()
                            }
                            supabase.table('usuarios').insert(nuevo_usuario).execute()
                            print(f"✅ Usuario creado en public.usuarios: {user.user.email}")
                    except Exception as e:
                        print(f"⚠️ Error verificando/creando usuario en public.usuarios: {e}")
                    
                    # Guardar usuario en sesión de Flask
                    session['user_id'] = user.user.id
                    session['user_email'] = user.user.email
                    session['user_nombre'] = user.user.user_metadata.get('nombre', email.split('@')[0])
                    
                    print(f"✅ Usuario logueado en sesión Flask: {user.user.email}")
                    
                    # Verificar si el usuario ya completó las preguntas desde la tabla usuarios
                    try:
                        print(f"🔍 Verificando preguntas completadas para usuario {user.user.id}")
                        user_response = supabase.table('usuarios').select('preguntas_completadas').eq('id', user.user.id).execute()
                        print(f"📊 Respuesta de la base de datos: {user_response.data}")
                        
                        if user_response.data:
                            preguntas_completadas = user_response.data[0].get('preguntas_completadas', 0)
                            print(f"✅ Usuario ya completó preguntas: {preguntas_completadas}")
                        else:
                            preguntas_completadas = 0
                            print(f"⚠️ Usuario no encontrado en la base de datos")
                    except Exception as e:
                        print(f"❌ Error verificando preguntas completadas: {e}")
                        preguntas_completadas = 0
                        
                    if not preguntas_completadas:
                        print(f"🔄 Redirigiendo a preguntas de usuario (preguntas_completadas={preguntas_completadas})")
                        return redirect(url_for("preguntas_usuario"))
                    
                    next_page = request.args.get('next')
                    if next_page:
                        print(f"🔄 Redirigiendo a: {next_page}")
                        return redirect(next_page)
                    else:
                        print(f"🔄 Redirigiendo a generar examen (preguntas_completadas={preguntas_completadas})")
                        return redirect(url_for('generar'))
                else:
                    print(f"❌ No se pudo autenticar usuario")
                    flash("Email o contraseña incorrectos")
                    
        except Exception as e:
            print(f"❌ Error en signin: {e}")
            print(f"🔍 Email intentado: {email}")
            print(f"🔍 Tipo de error: {type(e).__name__}")
            flash("Email o contraseña incorrectos")
            
        return render_template("login.html")
    
    return render_template("login.html")

@app.route("/auth/google")
def google_auth():
    """Iniciar autenticación con Google usando Supabase Auth"""
    try:
        if supabase:
            # Usar el método simple y directo de Supabase para OAuth
            response = supabase.auth.sign_in_with_oauth({
                "provider": "google",
                "options": {"redirect_to": "https://www.focusstudio.pw/auth/callback"}
            })
            
            print(f"🔄 Redirigiendo a Google OAuth...")
            print(f"🔍 URL de redirección: {response.url}")
            
            if response.url:
                return redirect(response.url)
            else:
                print(f"❌ No se obtuvo URL de redirección")
                flash("Error al iniciar autenticación con Google")
                return redirect(url_for('login'))
                
    except Exception as e:
        print(f"❌ Error iniciando Google OAuth: {e}")
        flash("Error al conectar con Google")
        return redirect(url_for('login'))

@app.route("/auth/callback")
def auth_callback():
    """Callback después de autenticación OAuth"""
    try:
        if supabase:
            # Obtener el código de autorización
            code = request.args.get('code')
            error = request.args.get('error')
            
            print(f"🔍 Callback recibido - Code: {code}, Error: {error}")
            
            if error:
                print(f"❌ Error en OAuth: {error}")
                flash(f"Error en la autenticación: {error}")
                return redirect(url_for('login'))
            
            if code:
                print(f"🔄 Procesando callback de Google OAuth...")
                
                # Intercambiar el code por una sesión en Supabase
                try:
                    supabase.auth.exchange_code_for_session({"auth_code": code})
                    user = supabase.auth.get_user()
                    
                    if user and hasattr(user, 'user'):
                        print(f"✅ Usuario obtenido después de OAuth: {user.user.email}")
                        
                        # Verificar si el usuario existe en public.usuarios, si no, crearlo
                        try:
                            usuario_check = supabase.table('usuarios').select('id').eq('id', user.user.id).execute()
                            if not usuario_check.data:
                                print(f"🔄 Usuario no existe en public.usuarios, creándolo...")
                                nuevo_usuario = {
                                    'id': user.user.id,
                                    'email': user.user.email,
                                    'nombre': user.user.user_metadata.get('nombre', user.user.email.split('@')[0]) if hasattr(user.user, 'user_metadata') and isinstance(user.user.user_metadata, dict) else user.user.email.split('@')[0],
                                    'fecha_registro': datetime.utcnow().isoformat(),
                                    'preguntas_completadas': 0,
                                    'total_examenes_rendidos': 0,
                                    'correctas_total': 0,
                                    'parciales_total': 0,
                                    'incorrectas_total': 0,
                                    'ultima_actividad': datetime.utcnow().isoformat()
                                }
                                supabase.table('usuarios').insert(nuevo_usuario).execute()
                                print(f"✅ Usuario creado en public.usuarios: {user.user.email}")
                        except Exception as e:
                            print(f"⚠️ Error verificando/creando usuario en public.usuarios: {e}")
                        
                        # Guardar usuario en sesión de Flask
                        session['user_id'] = user.user.id
                        session['user_email'] = user.user.email
                        
                        # Obtener nombre del user_metadata o usar email como fallback
                        if hasattr(user.user, 'user_metadata') and isinstance(user.user.user_metadata, dict):
                            session['user_nombre'] = user.user.user_metadata.get('nombre', user.user.email.split('@')[0])
                        else:
                            session['user_nombre'] = user.user.email.split('@')[0]
                        
                        print(f"✅ Usuario autenticado exitosamente: {user.user.email}")
                        
                        # Verificar si el usuario ya completó las preguntas
                        try:
                            user_response = supabase.table('usuarios').select('preguntas_completadas').eq('id', user.user.id).execute()
                            preguntas_completadas = user_response.data[0].get('preguntas_completadas', 0) if user_response.data else 0
                            
                            if not preguntas_completadas:
                                print(f"🔄 Redirigiendo a preguntas de usuario")
                                return redirect(url_for("preguntas_usuario"))
                            
                            print(f"🔄 Redirigiendo a generar examen")
                            return redirect(url_for('generar'))
                            
                        except Exception as e:
                            print(f"⚠️ Error verificando preguntas completadas: {e}")
                            return redirect(url_for("preguntas_usuario"))
                    else:
                        print(f"❌ No se pudo obtener usuario válido")
                        flash("Error en la autenticación con Google")
                        return redirect(url_for('login'))
                        
                except Exception as e:
                    print(f"❌ Error intercambiando código por sesión: {e}")
                    flash("Error en la autenticación con Google")
                    return redirect(url_for('login'))
            else:
                print(f"❌ No se recibió código de autorización")
                flash("Error en la autenticación con Google")
                return redirect(url_for('login'))
                
    except Exception as e:
        print(f"❌ Error en auth callback: {e}")
        flash("Error en la autenticación")
        return redirect(url_for('login'))

@app.route("/auth/logout")
def auth_logout():
    """Cerrar sesión con Supabase Auth"""
    try:
        if supabase:
            # Cerrar sesión en Supabase
            supabase.auth.sign_out()
            print(f"✅ Usuario cerró sesión en Supabase")
            
    except Exception as e:
        print(f"❌ Error en logout: {e}")
    
    # Limpiar sesión de Flask
    session.clear()
    print(f"✅ Sesión de Flask limpiada")
    return redirect(url_for("index"))

@app.route("/registro", methods=["GET", "POST"])
def registro():
    # Redirigir a la nueva ruta de Supabase Auth
    return redirect(url_for('signup'))

@app.route("/login", methods=["GET", "POST"])
def login():
    # Redirigir a la nueva ruta de Supabase Auth
    return redirect(url_for('signin'))

@app.route("/preguntas-usuario", methods=["GET", "POST"])
def preguntas_usuario():
    # Verificar autenticación simple
    print(f"🔍 Verificando autenticación en preguntas_usuario")
    print(f"🔍 Session data: {session}")
    print(f"🔍 is_authenticated(): {is_authenticated()}")
    
    if not is_authenticated():
        print(f"❌ Usuario no autenticado, redirigiendo a login")
        return redirect(url_for('login'))
    
    if request.method == "POST":
        como_nos_conociste = request.form.get("como_nos_conociste")
        uso_plataforma = request.form.get("uso_plataforma")

        try:
            if supabase:
                # Obtener usuario actual
                current_user = get_current_user()
                
                # Actualizar la tabla usuarios en Supabase
                print(f"🔄 Actualizando preguntas completadas para usuario {current_user['id']}")
                update_response = supabase.table('usuarios').update({
                    'como_nos_conociste': como_nos_conociste,
                    'plataforma_uso': uso_plataforma,
                    'preguntas_completadas': 1,
                    'ultima_actividad': datetime.utcnow().isoformat()
                }).eq('id', current_user['id']).execute()
                
                print(f"📊 Respuesta de actualización: {update_response.data}")
                
                if update_response.data:
                    print(f"✅ Información actualizada para usuario {current_user['email']}: como_nos_conociste={como_nos_conociste}, plataforma_uso={uso_plataforma}")
                    flash("Información guardada exitosamente!")
                    return redirect(url_for('generar'))
                else:
                    print(f"❌ No se pudo actualizar información para usuario {current_user['email']}")
                    flash("Error al guardar información. Intenta de nuevo.")
                    return render_template("preguntas_usuario.html")

        except Exception as e:
            print(f"Error guardando preguntas: {e}")
            flash("Error al guardar información. Intenta de nuevo.")

        return render_template("preguntas_usuario.html")

    return render_template("preguntas_usuario.html")

@app.route("/logout")
def logout():
    # Verificar autenticación simple
    if not is_authenticated():
        return redirect(url_for('login'))
    
    try:
        if supabase:
            current_user = get_current_user()
            # Invalidar sesión en Supabase (más seguro)
            try:
                supabase.auth.sign_out()
                print(f"✅ Sesión de Supabase invalidada para usuario {current_user['email']}")
            except Exception as e:
                print(f"⚠️ Error invalidando sesión de Supabase: {e}")
            
            # Log de la actividad
            log_data = {
                'usuario_id': current_user['id'],
                'tipo_actividad': 'logout',
                'fecha_actividad': datetime.utcnow().isoformat(),
                'detalles': {'accion': 'Usuario cerró sesión'},
                'ip_address': request.remote_addr
            }
            supabase.table('logs_actividad').insert(log_data).execute()
    except Exception as e:
        print(f"Error logging logout: {e}")

    # Limpiar sesión de Flask y invalidar cookie
    session.clear()
    response = redirect(url_for("index"))
    response.delete_cookie('session')  # Eliminar cookie de sesión
    return response

@app.route("/eliminar-cuenta", methods=["GET", "POST"])
def eliminar_cuenta():
    """Eliminar la cuenta del usuario y todos sus datos"""
    if not is_authenticated():
        return redirect(url_for('login'))
    
    if request.method == "POST":
        try:
            current_user = get_current_user()
            if not current_user:
                flash("Error: Usuario no encontrado")
                return redirect(url_for('perfil'))
            
            # Confirmar eliminación
            confirmacion = request.form.get('confirmacion', '').strip().lower()
            if confirmacion != 'eliminar':
                flash("Error: Debes escribir 'ELIMINAR' para confirmar la eliminación de tu cuenta")
                return render_template('eliminar_cuenta.html')
            
            print(f"🗑️ Iniciando eliminación de cuenta para usuario: {current_user['email']}")
            
            # Eliminar datos relacionados en Supabase
            if supabase:
                try:
                    # Eliminar exámenes del usuario
                    examenes_response = supabase.table('examenes').delete().eq('usuario_id', current_user['id']).execute()
                    print(f"✅ Exámenes eliminados: {len(examenes_response.data) if examenes_response.data else 0}")
                    
                    # Eliminar carpetas del usuario
                    carpetas_response = supabase.table('carpetas').delete().eq('usuario_id', current_user['id']).execute()
                    print(f"✅ Carpetas eliminadas: {len(carpetas_response.data) if carpetas_response.data else 0}")
                    
                    # Eliminar planificaciones del usuario
                    planificaciones_response = supabase.table('planificaciones').delete().eq('usuario_id', current_user['id']).execute()
                    print(f"✅ Planificaciones eliminadas: {len(planificaciones_response.data) if planificaciones_response.data else 0}")
                    
                    # Eliminar logs de actividad del usuario
                    logs_response = supabase.table('logs_actividad').delete().eq('usuario_id', current_user['id']).execute()
                    print(f"✅ Logs eliminados: {len(logs_response.data) if logs_response.data else 0}")
                    
                    # Eliminar el usuario de la tabla usuarios
                    usuario_response = supabase.table('usuarios').delete().eq('id', current_user['id']).execute()
                    print(f"✅ Usuario eliminado de tabla usuarios: {len(usuario_response.data) if usuario_response.data else 0}")
                    
                    # Nota: No eliminamos el usuario de Supabase Auth directamente
                    # ya que requiere permisos especiales. Los datos personales del usuario
                    # han sido eliminados de todas nuestras tablas, lo que efectivamente
                    # "elimina" su cuenta desde la perspectiva de la aplicación.
                    print(f"ℹ️ Datos personales eliminados. Usuario de Auth permanece inactivo.")
                    
                except Exception as e:
                    print(f"❌ Error eliminando datos de Supabase: {e}")
                    # Si ya eliminamos los datos principales, continuamos
                    # El usuario de Auth se eliminará en el paso siguiente
                    print(f"ℹ️ Continuando con eliminación de Auth...")
            
            # Limpiar sesión
            session.clear()
            
            print(f"✅ Cuenta eliminada exitosamente para: {current_user['email']}")
            flash("✅ Tu cuenta ha sido eliminada exitosamente. Todos tus datos personales (exámenes, carpetas, planificaciones) han sido borrados permanentemente. Tu sesión ha sido cerrada.")
            return redirect(url_for('index'))
            
        except Exception as e:
            print(f"❌ Error en eliminar_cuenta: {e}")
            flash("Error interno al eliminar la cuenta. Contacta al administrador.")
            return render_template('eliminar_cuenta.html')
    
    return render_template('eliminar_cuenta.html')

@app.route("/perfil")
def perfil():
    # Verificar autenticación simple
    if not is_authenticated():
        return redirect(url_for('login'))
    
    return render_template("perfil.html")

@app.route("/generar", methods=["GET", "POST"])
def generar():
    # Verificar autenticación simple
    if not is_authenticated():
        return redirect(url_for('login'))
    if request.method == "GET":
        try:
            if supabase:
                # Obtener carpetas del usuario para el selector
                carpetas_response = supabase.table('carpetas').select('id, nombre, color').eq('usuario_id', session.get('user_id')).order('nombre').execute()
                carpetas = carpetas_response.data if carpetas_response.data else []
            else:
                carpetas = []
        except Exception as e:
            print(f"⚠️ Error obteniendo carpetas: {e}")
            carpetas = []
        
        return render_template("generar.html", carpetas=carpetas)

    # Limpiar datos temporales de la sesión antes de generar un nuevo examen
    session.pop("preguntas", None)
    session.pop("respuestas", None)
    session.pop("pregunta_times", None)
    session.pop("start_time", None)
    session.pop("last_question_time", None)

    nivel = request.form["nivel"]
    cantidad = int(request.form["cantidad"])
    formato = request.form["formato"]
    archivo = request.files.get("archivo")
    tema = request.form.get("tema")
    cantidad_opciones = request.form.get("cantidad_opciones", "4")
    instrucciones_desarrollo = request.form.get("instrucciones_desarrollo", "")
    instrucciones_vf = request.form.get("instrucciones_vf", "")
    temas_math = request.form.getlist("temas")
    tema_personalizado = request.form.get("tema_personalizado", "")

    # --- NUEVO: Ejercicios matemáticos ---
    if formato == "ejercicios matematicos":
        temas = temas_math.copy()
        if tema_personalizado:
            temas.append(tema_personalizado)
        if not temas:
            return "Debes seleccionar al menos un tema de matemática."
        ejercicios = []
        for i in range(cantidad):
            # Generar enunciado con GPT-4
            prompt = (
                f"Generá un ejercicio matemático de nivel {nivel} sobre el tema '{temas[i % len(temas)]}'. "
                "El ejercicio debe tener UNA sola consigna, ser claro, concreto y estar expresado como una expresión matemática o pregunta directa, NO como un problema con partes a) y b). "
                "Incluí la expresión matemática principal entre corchetes al final, por ejemplo: [expresión]. No incluyas la solución ni la respuesta."
            )
            try:
                print(f"🔍 [MATEMÁTICAS] Generando ejercicio {i+1}/{cantidad}...")
                print(f"🔍 [MATEMÁTICAS] Longitud del prompt: {len(prompt)} caracteres")
                
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": "Sos un generador de ejercicios matemáticos para exámenes."},
                        {"role": "user", "content": prompt}
                    ],
                    max_completion_tokens=120,
                    timeout=30
                )
                
                # Log de tokens utilizados
                if hasattr(response, 'usage'):
                    print(f"🔍 [MATEMÁTICAS] Tokens utilizados:")
                    print(f"   📥 Input tokens: {response.usage.prompt_tokens}")
                    print(f"   📤 Output tokens: {response.usage.completion_tokens}")
                    print(f"   📊 Total tokens: {response.usage.total_tokens}")
                
                enunciado_gpt = response.choices[0].message.content.strip()
                print(f"🔍 [MATEMÁTICAS] Ejercicio generado: {len(enunciado_gpt)} caracteres")
            except Exception as e:
                enunciado_gpt = f"[Error al generar enunciado: {str(e)}]"
            # Extraer expresión entre corchetes
            match = re.search(r"\[(.*?)\]", enunciado_gpt)
            expresion = match.group(1) if match else ""
            # Limpiar delimitadores LaTeX si existen
            expresion = expresion.replace('\\(', '').replace('\\)', '').strip()
            enunciado = enunciado_gpt.replace(f'[{match.group(1)}]', '').strip() if match else enunciado_gpt
            # Obtener imagen y solución con Wolfram usando la expresión
            try:
                url = "https://api.wolframalpha.com/v2/query"
                params = {
                    "input": expresion,
                    "appid": app_id,
                    "format": "image,plaintext"
                }
                resp = requests.get(url, params=params, timeout=30)
                solucion = ""
                pods = []
                img_enunciado = None
                if resp.status_code == 200:
                    import xml.etree.ElementTree as ET
                    root = ET.fromstring(resp.text)
                    for pod in root.findall(".//pod"):
                        pod_title = pod.attrib.get("title", "")
                        subpod = pod.find("subpod")
                        pod_plaintext = subpod.findtext("plaintext") if subpod is not None else None
                        pod_img = None
                        if subpod is not None:
                            img_tag = subpod.find("img")
                            if img_tag is not None:
                                pod_img = img_tag.attrib.get("src")
                        # Guardar la imagen del enunciado (primer pod Input)
                        if not img_enunciado and pod_title.lower() in ["input", "entrada"] and pod_img:
                            img_enunciado = pod_img
                        if pod_title.lower() in ["result", "resultado", "solution", "solución"] and pod_plaintext:
                            solucion = pod_plaintext
                        if pod_plaintext or pod_img:
                            pods.append({"title": pod_title, "plaintext": pod_plaintext, "img": pod_img})
                else:
                    solucion = "[Error al consultar Wolfram Alpha]"
                    img_enunciado = None
            except Exception as e:
                solucion = f"[Error al consultar Wolfram: {str(e)}]"
                pods = []
                img_enunciado = None
            ejercicios.append({
                "enunciado": enunciado,
                "expresion": expresion,
                "img_enunciado": img_enunciado,
                "solucion": solucion,
                "pods": pods,
                "respuesta_usuario": ""
            })
        session["ejercicios_matematicos"] = ejercicios
        session["start_time"] = time.time()
        return redirect(url_for("examen_matematico", numero=0))
    # --- FIN NUEVO ---

    texto = ""
    if archivo and archivo.filename:
        if archivo.filename.endswith(".txt"):
            print(f"\n--- PROCESANDO ARCHIVO TXT: {archivo.filename} ---")
            texto = archivo.read().decode("utf-8")
            print(f"Total de caracteres: {len(texto)}")
            print(f"Primeros 500 caracteres: {texto[:500]}...")
            print("--- FIN ARCHIVO TXT ---\n")
        elif archivo.filename.endswith(".pdf"):
            pdf_stream = BytesIO(archivo.read())
            reader = PyPDF2.PdfReader(pdf_stream)
            print(f"\n--- PROCESANDO PDF: {archivo.filename} ---")
            print(f"Total de páginas: {len(reader.pages)}")
            
            texto = ""
            max_pages = min(len(reader.pages), 7)  # 🎯 LÍMITE: Solo las primeras 7 páginas
            print(f"Procesando primeras {max_pages} páginas de {len(reader.pages)} totales")
            
            for i in range(max_pages):
                try:
                    page_text = reader.pages[i].extract_text()
                    if page_text and page_text.strip():
                        texto += f"\n--- PÁGINA {i+1} ---\n{page_text.strip()}\n"
                        print(f"Página {i+1}: {len(page_text)} caracteres extraídos")
                    else:
                        print(f"Página {i+1}: Sin texto extraído (página vacía o imagen)")
                except Exception as e:
                    print(f"Página {i+1}: Error al extraer texto - {str(e)}")
            
            if len(reader.pages) > max_pages:
                print(f"⚠️  NOTA: Solo se procesaron las primeras {max_pages} páginas de {len(reader.pages)}")
                print(f"💡 Para procesar más páginas, considera dividir el PDF en archivos más pequeños")
            
            # Si no se extrajo texto, intentar métodos alternativos
            if not texto.strip():
                print("⚠️  ADVERTENCIA: No se pudo extraer texto del PDF")
                print("💡 Posibles causas:")
                print("   - PDF escaneado (solo imágenes)")
                print("   - PDF con protección DRM")
                print("   - PDF con fuentes especiales")
                print("   - PDF con layout complejo")
                print("💡 Soluciones:")
                print("   - Usar un PDF con texto seleccionable")
                print("   - Convertir el PDF a texto primero")
                print("   - Usar un archivo TXT o DOCX en su lugar")
            
            print(f"\n--- RESUMEN PDF ---")
            print(f"Archivo: {archivo.filename}")
            print(f"Páginas procesadas: {len(reader.pages)}")
            print(f"Total de caracteres extraídos: {len(texto)}")
            if texto.strip():
                print(f"Primeros 500 caracteres: {texto[:500]}...")
            else:
                print("❌ NO SE EXTRAJO TEXTO DEL PDF")
            print("--- FIN RESUMEN PDF ---\n")
        elif archivo.filename.endswith(".docx"):
            print(f"\n--- PROCESANDO ARCHIVO DOCX: {archivo.filename} ---")
            docx_stream = BytesIO(archivo.read())
            doc = docx.Document(docx_stream)
            print(f"Total de párrafos: {len(doc.paragraphs)}")
            
            texto = ""
            max_paragraphs = min(len(doc.paragraphs), 50)  # 🎯 LÍMITE: Solo los primeros 50 párrafos (equivalente a ~7 páginas)
            print(f"Procesando primeros {max_paragraphs} párrafos de {len(doc.paragraphs)} totales")
            
            for i in range(max_paragraphs):
                if doc.paragraphs[i].text.strip():
                    texto += doc.paragraphs[i].text + "\n"
                    print(f"Párrafo {i+1}: {len(doc.paragraphs[i].text)} caracteres")
            
            if len(doc.paragraphs) > max_paragraphs:
                print(f"⚠️  NOTA: Solo se procesaron los primeros {max_paragraphs} párrafos de {len(doc.paragraphs)}")
                print(f"💡 Para procesar más contenido, considera dividir el documento en archivos más pequeños")
            
            print(f"\n--- RESUMEN DOCX ---")
            print(f"Archivo: {archivo.filename}")
            print(f"Párrafos procesados: {len(doc.paragraphs)}")
            print(f"Total de caracteres extraídos: {len(texto)}")
            print(f"Primeros 500 caracteres: {texto[:500]}...")
            print("--- FIN RESUMEN DOCX ---\n")
        print("\n--- TEXTO EXTRAÍDO DEL ARCHIVO (GENERADOR) ---\n", texto, "\n--- FIN TEXTO EXTRAÍDO ---\n")
        
        # Validar que se extrajo texto del archivo
        if archivo and archivo.filename and not texto.strip():
            if archivo.filename.endswith(".pdf"):
                return render_template("generar.html", mensaje_error="No se pudo extraer texto del PDF. Posibles causas: PDF escaneado, con protección DRM, o con fuentes especiales. Intenta con un PDF que tenga texto seleccionable o usa un archivo TXT/DOCX.")
            else:
                return render_template("generar.html", mensaje_error="No se pudo extraer texto del archivo. Verifica que el archivo contenga texto legible.")
    elif tema:
        texto = f"Tema: {tema}."
    else:
        return "Debe ingresar un tema o subir un archivo."

    if formato == "multiple choice":
        prompt = (
            f"Generá {cantidad} preguntas de examen en formato opción múltiple, todas directamente relacionadas con el siguiente tema, para nivel {nivel}. "
            f"Cada pregunta debe comenzar con 'Enunciado X: ...', incluir {cantidad_opciones} opciones (a, b, c, d" + (", e" if cantidad_opciones == "5" else "") + (", c" if cantidad_opciones == "3" else "") + ") en líneas separadas, asegurando que solo una opción sea correcta y las otras sean plausibles y relacionadas con el tema (no obvias ni irrelevantes). "
            "Al final de cada pregunta, escribí: Respuesta: x. Evitá preguntas demasiado generales o de sentido común."
        )
    elif formato == "verdadero o falso":
        prompt = (
            f"Generá {cantidad} preguntas en formato verdadero o falso, todas directamente relacionadas con el siguiente tema, para nivel {nivel}. "
            "Cada pregunta debe comenzar con 'Enunciado X: Seleccionar verdadero o falso: ...', ser conceptualmente profunda y no trivial, y terminar con 'Respuesta: Verdadero' o 'Respuesta: Falso'. "
            "Evitá afirmaciones obvias o que no requieran conocimiento del tema. "
            + (f"Instrucciones adicionales: {instrucciones_vf}" if instrucciones_vf else "")
        )
    else:
        prompt = (
            f"Generá {cantidad} preguntas de examen abiertas para que el estudiante responda desarrollando, basadas en el siguiente tema, en orden aleatorio, para nivel {nivel}. "
            "Comenzá cada una con 'Enunciado X: ...'. No incluyas opciones ni respuesta. "
            + (f"Instrucciones adicionales: {instrucciones_desarrollo}" if instrucciones_desarrollo else "")
        )

    # LOG: Mostrar el prompt enviado a la IA
    print("\n--- PROMPT ENVIADO A LA IA ---\n", prompt + "\n\n" + texto, "\n--- FIN PROMPT ---\n")

    try:
        print(f"🔍 [GENERADOR] Enviando prompt a GPT-4o...")
        print(f"🔍 [GENERADOR] API Key presente: {'Sí' if os.getenv('OPENAI_API_KEY') else 'No'}")
        print(f"🔍 [GENERADOR] Longitud del prompt: {len(prompt)} caracteres")
        print(f"🔍 [GENERADOR] Longitud del texto: {len(texto)} caracteres")
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "Sos un generador de exámenes"},
                {"role": "user", "content": prompt + "\n\n" + texto}
            ],
            max_completion_tokens=3000,
            timeout=60
        )
        
        # Log de tokens utilizados
        if hasattr(response, 'usage'):
            print(f"🔍 [GENERADOR] Tokens utilizados:")
            print(f"   📥 Input tokens: {response.usage.prompt_tokens}")
            print(f"   📤 Output tokens: {response.usage.completion_tokens}")
            print(f"   📊 Total tokens: {response.usage.total_tokens}")
        
        print(f"✅ [GENERADOR] Respuesta recibida de GPT-4o")
        preguntas_raw = response.choices[0].message.content
        print(f"🔍 [GENERADOR] Respuesta recibida: {len(preguntas_raw)} caracteres")
        print(f"🔍 [DEBUG] Longitud de la respuesta: {len(preguntas_raw)}")

        # Log para depuración: ver qué devuelve la IA
        print("\n\n--- RESPUESTA IA RAW ---\n", preguntas_raw, "\n--- FIN RESPUESTA ---\n\n")

        if preguntas_raw:
            # Dividir por líneas y procesar manualmente
            lineas = preguntas_raw.strip().split('\n')
            bloques = []
            bloque_actual = []
            en_pregunta = False
            
            for linea in lineas:
                linea = linea.strip()
                if not linea:
                    continue
                
                # Detectar inicio de nueva pregunta
                if linea.startswith('**Enunciado') or linea.startswith('Enunciado'):
                    if bloque_actual:
                        bloques.append('\n'.join(bloque_actual))
                        bloque_actual = []
                    en_pregunta = True
                    bloque_actual.append(linea)
                elif en_pregunta:
                    bloque_actual.append(linea)
                    # Detectar fin de pregunta (cuando encontramos Respuesta:)
                    if linea.startswith('Respuesta:'):
                        en_pregunta = False
            
            # Agregar el último bloque
            if bloque_actual:
                bloques.append('\n'.join(bloque_actual))
            
            print(f"\n--- BLOQUES ENCONTRADOS: {len(bloques)} ---\n")
            for i, bloque in enumerate(bloques):
                print(f"Bloque {i+1}: {bloque[:100]}...")
        else:
            bloques = []
            print("\n--- NO SE ENCONTRARON BLOQUES ---\n")
        
        preguntas = []
        for bloque in bloques:
            lineas = bloque.strip().split("\n")
            enunciado = next((l for l in lineas if l.lower().startswith("enunciado")), lineas[0])
            opciones = []
            respuesta = "indefinida"
            # tipo = "desarrollo"  # El tipo ahora se fuerza según el formato seleccionado

            # Determinar qué opciones buscar según la cantidad configurada
            opciones_buscar = []
            if cantidad_opciones == "3":
                opciones_buscar = ["a)", "b)", "c)"]
            elif cantidad_opciones == "5":
                opciones_buscar = ["a)", "b)", "c)", "d)", "e)"]
            else:  # 4 opciones por defecto
                opciones_buscar = ["a)", "b)", "c)", "d)"]

            for l in lineas:
                l_strip = l.strip().lower()
                if l_strip.startswith("respuesta"):
                    raw_resp = l.split(":")[-1].strip().lower().rstrip('.')  # Remover punto al final
                    if raw_resp in ["verdadero", "falso"]:
                        respuesta = raw_resp
                    elif raw_resp in ["a", "b", "c", "d", "e"]:
                        respuesta = raw_resp
                if any(l_strip.startswith(op) for op in opciones_buscar):
                    opciones.append(l.strip())

            # Forzar el tipo según la selección del usuario
            if formato == "multiple choice":
                tipo = "multiple"
            elif formato == "verdadero o falso":
                tipo = "vf"
            else:
                tipo = "desarrollo"

            if tipo in ["multiple", "vf"] and respuesta == "indefinida":
                print(f"\n--- PREGUNTA DESCARTADA: {enunciado[:50]}... ---\n")
                continue

            preguntas.append({"enunciado": enunciado, "opciones": opciones, "respuesta": respuesta, "tipo": tipo, "tema": "General"})
            print(f"\n--- PREGUNTA AGREGADA: {enunciado[:50]}... (tipo: {tipo}, respuesta: {respuesta}) ---\n")

        # Validar que haya preguntas y que todos los enunciados sean válidos
        if not preguntas or any(not p["enunciado"].strip() for p in preguntas):
            mensaje_error = "No se pudieron generar preguntas válidas. Verificá el texto, el formato o intentá nuevamente."
            print(f"\n--- ERROR: {mensaje_error} ---\n")
            return render_template("generar.html", mensaje_error=mensaje_error)

        print(f"\n--- PREGUNTAS FINALES: {len(preguntas)} ---\n")
        session["preguntas"] = preguntas
        session["respuestas"] = ["" for _ in preguntas]
        session["start_time"] = time.time()
        session["pregunta_times"] = []
        session["last_question_time"] = time.time()
        
        # Guardar la carpeta seleccionada en la sesión
        carpeta_id = request.values.get("carpeta_id")
        if carpeta_id:
            session['carpeta_seleccionada'] = carpeta_id
            print(f"💾 Carpeta guardada en sesión ANTES de ir a pregunta: {carpeta_id}")
        else:
            session.pop('carpeta_seleccionada', None)
            print("🗑️ No hay carpeta seleccionada, limpiando sesión")
        
        # Guardar el título personalizado en la sesión
        titulo = request.values.get("titulo")
        if titulo:
            session['titulo_examen'] = titulo
            print(f"📝 Título guardado en sesión: {titulo}")
        else:
            session.pop('titulo_examen', None)
            print("🗑️ No hay título ingresado, limpiando sesión")

    except Exception as e:
        print(f"\n--- EXCEPCIÓN: {str(e)} ---\n")
        print(f"🔍 [DEBUG] Tipo de error: {type(e).__name__}")
        print(f"🔍 [DEBUG] Detalles del error: {str(e)}")
        import traceback
        traceback.print_exc()
        
        # Manejo específico para timeouts
        if "timeout" in str(e).lower() or "timed out" in str(e).lower():
            return render_template("generar.html", mensaje_error="La generación tardó demasiado tiempo en Heroku. Por favor, intenta con un archivo más pequeño, menos preguntas, o usa la versión local.")
        else:
            return render_template("generar.html", mensaje_error=f"Error al generar preguntas: {str(e)}. Por favor, intenta nuevamente.")

    # Resumen final de tokens utilizados
    print(f"\n📊 [RESUMEN GENERADOR] Función completada exitosamente")
    print(f"📊 [RESUMEN GENERADOR] Total de preguntas generadas: {len(preguntas)}")
    print(f"📊 [RESUMEN GENERADOR] Revisar logs anteriores para detalles de tokens por función")
    
    return redirect(url_for('pregunta', numero=0))

@app.route("/pregunta/<int:numero>", methods=["GET", "POST"])
def pregunta(numero):
    preguntas = session.get("preguntas", [])
    respuestas = session.get("respuestas", [])

    if numero >= len(preguntas):
        return redirect(url_for('resultado'))

    if request.method == "POST":
        respuesta_usuario = request.form.get("respuesta", "")
        respuestas[numero] = respuesta_usuario
        session["respuestas"] = respuestas

        now = time.time()
        duracion = now - session["last_question_time"]
        session["pregunta_times"].append(round(duracion, 2))
        session["last_question_time"] = now

        return redirect(url_for('pregunta', numero=numero + 1))

    pregunta = preguntas[numero]
    respuesta_actual = respuestas[numero]
    return render_template("examen.html", pregunta=pregunta, numero=numero + 1, total=len(preguntas), actual=numero, respuesta_actual=respuesta_actual)

@app.route("/resultado")
def resultado():
    # Verificar autenticación simple
    if not is_authenticated():
        return redirect(url_for('login'))
    preguntas = session.get("preguntas", [])
    respuestas = session.get("respuestas", [])
    tiempos = session.get("pregunta_times", [])

    feedbacks = []
    correctas = 0
    parciales = 0
    incorrectas = 0
    temas_fallidos = {}
    preguntas_falladas = []

    respuestas_texto_usuario = []
    respuestas_texto_correcta = []
    for i in range(len(preguntas)):
        pregunta = preguntas[i]
        respuesta_usuario = respuestas[i]
        texto_usuario = ""
        texto_correcta = ""
        if pregunta["tipo"] == "multiple":
            for op in pregunta["opciones"]:
                if respuesta_usuario and op.lower().startswith(f"{respuesta_usuario})"):
                    texto_usuario = op[2:].strip()
                if op.lower().startswith(f"{pregunta['respuesta']})"):
                    texto_correcta = op[2:].strip()
        respuestas_texto_usuario.append(texto_usuario)
        respuestas_texto_correcta.append(texto_correcta)

    for i in range(len(preguntas)):
        pregunta = preguntas[i]
        respuesta_usuario = respuestas[i]
        feedback = ""
        explicacion_ia = ""

        if pregunta["tipo"] in ["multiple", "vf"]:
            correcta = pregunta["respuesta"]
            texto_correcta = ""
            texto_usuario = ""
            if pregunta["tipo"] == "multiple":
                for op in pregunta["opciones"]:
                    if op.lower().startswith(f"{correcta})"):
                        texto_correcta = op[2:].strip()
                    if respuesta_usuario and op.lower().startswith(f"{respuesta_usuario})"):
                        texto_usuario = op[2:].strip()
            if respuesta_usuario == correcta:
                feedback = f"✔️ CORRECTA"
                correctas += 1
            else:
                # --- FEEDBACK IA BREVE ---
                try:
                    prompt_ia = (
                        f"Sos un profesor que explica brevemente por qué una respuesta es correcta en un examen. "
                        f"Esta es la pregunta de examen: {pregunta['enunciado']}\n"
                        f"El alumno respondió: {respuesta_usuario}\n"
                        f"La respuesta correcta es: {correcta}\n"
                        f"Explicá en 1-2 frases, de forma breve y clara, por qué la respuesta correcta es la que corresponde. No repitas el enunciado completo."
                    )
                    print(f"🔍 [EXPLICACIÓN] Generando explicación para pregunta {i+1}...")
                    print(f"🔍 [EXPLICACIÓN] Longitud del prompt: {len(prompt_ia)} caracteres")
                    
                    response = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[
                            {"role": "system", "content": "Sos un profesor que explica brevemente por qué una respuesta es correcta en un examen."},
                            {"role": "user", "content": prompt_ia}
                        ],
                        max_completion_tokens=120,
                        timeout=30
                    )
                    
                    # Log de tokens utilizados
                    if hasattr(response, 'usage'):
                        print(f"🔍 [EXPLICACIÓN] Tokens utilizados:")
                        print(f"   📥 Input tokens: {response.usage.prompt_tokens}")
                        print(f"   📤 Output tokens: {response.usage.completion_tokens}")
                        print(f"   📊 Total tokens: {response.usage.total_tokens}")
                    
                    explicacion_ia = response.choices[0].message.content.strip()
                    print(f"🔍 [EXPLICACIÓN] Explicación generada: {len(explicacion_ia)} caracteres")
                except Exception as e:
                    explicacion_ia = "(No se pudo generar explicación IA)"
                if pregunta["tipo"] == "multiple":
                    feedback = f"❌ INCORRECTA.\nTu respuesta fue '{respuesta_usuario}': \"{texto_usuario}\"\nLa correcta era '{correcta}': \"{texto_correcta}\"\n<b>Por qué: </b>{explicacion_ia}"
                else:
                    feedback = f"❌ INCORRECTA. Tu respuesta fue '{respuesta_usuario}', la correcta era '{correcta}'.\n<b>Por qué: </b>{explicacion_ia}"
                incorrectas += 1
                temas_fallidos[pregunta["tema"]] = temas_fallidos.get(pregunta["tema"], 0) + 1
                preguntas_falladas.append({
                    "enunciado": pregunta["enunciado"],
                    "respuesta_usuario": respuesta_usuario,
                    "respuesta_correcta": correcta
                })
        else:
            prompt = (
                f"Pregunta: {pregunta['enunciado']}\n"
                f"Respuesta del alumno: {respuesta_usuario}\n"
                "Evaluá si la respuesta es correcta, incorrecta o parcialmente correcta y explicá brevemente por qué. "
                "Al final, decí solo CORRECTA, INCORRECTA o PARCIALMENTE CORRECTA."
            )
            try:
                print(f"🔍 [FEEDBACK] Generando feedback para pregunta {i+1}...")
                print(f"🔍 [FEEDBACK] Longitud del prompt: {len(prompt)} caracteres")
                
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": "Sos un corrector de exámenes"},
                        {"role": "user", "content": prompt}
                    ],
                    max_completion_tokens=500,
                    timeout=45
                )
                
                # Log de tokens utilizados
                if hasattr(response, 'usage'):
                    print(f"🔍 [FEEDBACK] Tokens utilizados:")
                    print(f"   📥 Input tokens: {response.usage.prompt_tokens}")
                    print(f"   📤 Output tokens: {response.usage.completion_tokens}")
                    print(f"   📊 Total tokens: {response.usage.total_tokens}")
                
                feedback_raw = response.choices[0].message.content
                print(f"🔍 [FEEDBACK] Feedback generado: {len(feedback_raw)} caracteres")

                if feedback_raw:
                    f_lower = feedback_raw.lower()
                else:
                    f_lower = ""

                if "parcialmente correcta" in f_lower:
                    feedback = f"⚠️ PARCIALMENTE CORRECTA\n{feedback_raw}"
                    parciales += 1
                    temas_fallidos[pregunta["tema"]] = temas_fallidos.get(pregunta["tema"], 0) + 1
                    preguntas_falladas.append({
                        "enunciado": pregunta["enunciado"],
                        "respuesta_usuario": respuesta_usuario,
                        "respuesta_correcta": "(respuesta abierta)"
                    })
                elif "incorrecta" in f_lower:
                    feedback = f"❌ INCORRECTA\n{feedback_raw}"
                    incorrectas += 1
                    temas_fallidos[pregunta["tema"]] = temas_fallidos.get(pregunta["tema"], 0) + 1
                    preguntas_falladas.append({
                        "enunciado": pregunta["enunciado"],
                        "respuesta_usuario": respuesta_usuario,
                        "respuesta_correcta": "(respuesta abierta)"
                    })
                elif "correcta" in f_lower:
                    feedback = f"✔️ CORRECTA\n{feedback_raw}"
                    correctas += 1
                else:
                    feedback = f"⚠️ No se pudo clasificar la respuesta\n{feedback_raw}"
            except Exception as e:
                feedback = f"Error al corregir: {str(e)}"

        feedbacks.append(feedback)

    total = len(preguntas)
    nota = round((correctas + parciales * 0.5) / total * 10, 2)

    # --- FEEDBACK GENERAL IA ---
    if preguntas_falladas:
        try:
            prompt_general = (
                "Sos un tutor experto en ayudar a estudiantes a mejorar en exámenes. "
                "Te paso una lista de preguntas que el estudiante respondió incorrectamente o parcialmente, junto con su respuesta y la respuesta correcta. "
                "En base a estos errores, respondé en segunda persona y comenzá tu respuesta con 'Te recomendamos enfocarte en...'. "
                "Sé concreto, breve (2-3 frases) y no repitas el enunciado de las preguntas.\n\nPreguntas falladas:\n"
            )
            for pf in preguntas_falladas:
                prompt_general += f"- Enunciado: {pf['enunciado']}\n  Respuesta del alumno: {pf['respuesta_usuario']}\n  Respuesta correcta: {pf['respuesta_correcta']}\n"
            print(f"🔍 [FEEDBACK GENERAL] Generando feedback general...")
            print(f"🔍 [FEEDBACK GENERAL] Longitud del prompt: {len(prompt_general)} caracteres")
            
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "Sos un tutor experto en ayudar a estudiantes a mejorar en exámenes."},
                    {"role": "user", "content": prompt_general}
                ],
                max_completion_tokens=200,
                timeout=45
            )
            
            # Log de tokens utilizados
            if hasattr(response, 'usage'):
                print(f"🔍 [FEEDBACK GENERAL] Tokens utilizados:")
                print(f"   📥 Input tokens: {response.usage.prompt_tokens}")
                print(f"   📤 Output tokens: {response.usage.completion_tokens}")
                print(f"   📊 Total tokens: {response.usage.total_tokens}")
            
            feedback_general = response.choices[0].message.content.strip()
            print(f"🔍 [FEEDBACK GENERAL] Feedback general generado: {len(feedback_general)} caracteres")
        except Exception as e:
            feedback_general = "(No se pudo generar feedback personalizado)"
    else:
        feedback_general = "¡Excelente! No se detectaron temas con errores frecuentes."

    resumen = {
        "correctas": correctas,
        "parciales": parciales,
        "incorrectas": incorrectas,
        "total": total,
        "nota": nota,
        "tiempo_total": round(time.time() - session["start_time"], 2),
        "tiempos_por_pregunta": tiempos,
        "feedback_general": feedback_general
    }

    # GUARDAR EN BASE DE DATOS SUPABASE
    if is_authenticated() and supabase:
        try:
            current_user = get_current_user()
            print(f"\n🔍 INTENTANDO GUARDAR EXAMEN EN SUPABASE...")
            print(f"Usuario: {current_user['email']} (ID: {current_user['id']})")
            print(f"Nota: {nota}/10")
            print(f"Tiempo: {resumen['tiempo_total']}s")
            
            # Verificar que las tablas existan y crear usuario si no existe
            try:
                # Verificar tabla examenes
                examenes_check = supabase.table('examenes').select('id').limit(1).execute()
                print(f"✅ Tabla 'examenes' existe")
                
                # Verificar tabla preguntas_examen
                preguntas_check = supabase.table('preguntas_examen').select('id').limit(1).execute()
                print(f"✅ Tabla 'preguntas_examen' existe")
                
                # Verificar si el usuario existe en public.usuarios, si no, crearlo
                try:
                    usuario_check = supabase.table('usuarios').select('id').eq('id', current_user['id']).execute()
                    if not usuario_check.data:
                        print(f"🔄 Usuario no existe en public.usuarios, creándolo...")
                        nuevo_usuario = {
                            'id': current_user['id'],
                            'email': current_user['email'],
                            'nombre': current_user['nombre'],
                            'fecha_registro': datetime.utcnow().isoformat(),
                            'preguntas_completadas': 0,
                            'total_examenes_rendidos': 0,
                            'correctas_total': 0,
                            'parciales_total': 0,
                            'incorrectas_total': 0,
                            'ultima_actividad': datetime.utcnow().isoformat()
                        }
                        supabase.table('usuarios').insert(nuevo_usuario).execute()
                        print(f"✅ Usuario creado en public.usuarios: {current_user['email']}")
                    else:
                        print(f"✅ Usuario encontrado en tabla 'usuarios'")
                except Exception as e:
                    print(f"⚠️ Error verificando/creando usuario en public.usuarios: {e}")
                    return render_template("resultado_abierto.html", respuestas=respuestas, preguntas=preguntas, feedbacks=feedbacks, resumen=resumen, respuestas_texto_usuario=respuestas_texto_usuario, respuestas_texto_correcta=respuestas_texto_correcta)
                
            except Exception as e:
                print(f"❌ Error verificando tablas: {e}")
                return render_template("resultado_abierto.html", respuestas=respuestas, preguntas=preguntas, feedbacks=feedbacks, resumen=resumen, respuestas_texto_usuario=respuestas_texto_usuario, respuestas_texto_correcta=respuestas_texto_correcta)
            
            # Obtener carpeta seleccionada desde la sesión (ya guardada anteriormente)
            carpeta_id = session.get('carpeta_seleccionada')
            
            # Guardar examen principal
            examen_data = {
                'usuario_id': current_user['id'],
                'titulo': session.get('titulo_examen', f'Examen de {preguntas[0].get("tema", "General")}'),
                'materia': preguntas[0].get("tema", "General"),
                'fecha_creacion': datetime.utcnow().isoformat(),
                'fecha_rendido': datetime.utcnow().isoformat(),
                'preguntas': json.dumps([p['enunciado'] for p in preguntas]),
                'respuestas': json.dumps(respuestas),
                'nota': nota,
                'tiempo_duracion': int(float(resumen["tiempo_total"])),
                'estado': 'rendido',
                'tiempo_total_segundos': int(float(resumen["tiempo_total"])),
                # Agregar métricas detalladas
                'correctas': correctas,
                'parciales': parciales,
                'incorrectas': incorrectas,
                'total_preguntas': total,
                # Agregar feedback general
                'feedback_general': feedback_general,
                # Agregar carpeta si fue seleccionada (usar la de la sesión)
                'carpeta_id': session.get('carpeta_seleccionada') if session.get('carpeta_seleccionada') else None
            }
            
            print(f"📊 Datos del examen: {examen_data}")
            
            # Limpiar datos de la sesión después de guardar
            session.pop('titulo_examen', None)
            session.pop('carpeta_seleccionada', None)
            
            examen_response = supabase.table('examenes').insert(examen_data).execute()
            
            if examen_response.data:
                examen_id = examen_response.data[0]['id']
                print(f"✅ Examen guardado con ID: {examen_id}")
                
                # Guardar cada pregunta individual
                print(f"📝 Guardando {len(preguntas)} preguntas...")
                for i, pregunta in enumerate(preguntas):
                    pregunta_data = {
                        'examen_id': examen_id,
                        'pregunta': pregunta['enunciado'],
                        'opciones': json.dumps(pregunta.get('opciones', [])),
                        'respuesta_usuario': respuestas[i],
                        'respuesta_correcta': pregunta['respuesta'],
                        'es_correcta': respuestas[i] == pregunta['respuesta'],
                        'es_parcial': False,  # Por defecto, puedes ajustar según tu lógica
                        'feedback': '',  # Por defecto vacío
                        'orden': i + 1,
                        'tiempo_respuesta': 0  # Por defecto 0, puedes ajustar si tienes este dato
                    }
                    
                    print(f"  Pregunta {i+1}: {pregunta['enunciado'][:50]}...")
                    pregunta_response = supabase.table('preguntas_examen').insert(pregunta_data).execute()
                    
                    if pregunta_response.data:
                        print(f"    ✅ Pregunta {i+1} guardada")
                    else:
                        print(f"    ❌ Error guardando pregunta {i+1}")
                
                # Actualizar estadísticas del usuario
                try:
                    # Obtener el usuario actual
                    user_response = supabase.table('usuarios').select('total_examenes_rendidos, correctas_total, parciales_total, incorrectas_total').eq('id', current_user['id']).execute()
                    if user_response.data:
                        user_data = user_response.data[0]
                        total_actual = user_data.get('total_examenes_rendidos', 0)
                        correctas_actual = user_data.get('correctas_total', 0)
                        parciales_actual = user_data.get('parciales_total', 0)
                        incorrectas_actual = user_data.get('incorrectas_total', 0)
                        
                        nuevo_total = total_actual + 1
                        nuevo_correctas = correctas_actual + correctas
                        nuevo_parciales = parciales_actual + parciales
                        nuevo_incorrectas = incorrectas_actual + incorrectas
                        
                        # Actualizar el contador
                        supabase.table('usuarios').update({
                            'total_examenes_rendidos': nuevo_total,
                            'correctas_total': nuevo_correctas,
                            'parciales_total': nuevo_parciales,
                            'incorrectas_total': nuevo_incorrectas,
                            'ultima_actividad': datetime.utcnow().isoformat()
                        }).eq('id', current_user['id']).execute()
                        
                        print(f"✅ Estadísticas actualizadas: {nuevo_total} exámenes, {nuevo_correctas} correctas, {nuevo_parciales} parciales, {nuevo_incorrectas} incorrectas")
                except Exception as e:
                    print(f"⚠️ Error actualizando estadísticas: {e}")
                
                # Guardar estadísticas diarias
                try:
                    fecha_hoy = datetime.utcnow().date().isoformat()
                    
                    # Verificar si ya existen estadísticas para hoy
                    stats_response = supabase.table('estadisticas_usuarios').select('*').eq('usuario_id', current_user['id']).eq('fecha_estadistica', fecha_hoy).execute()
                    
                    if stats_response.data:
                        # Actualizar estadísticas existentes
                        stats_id = stats_response.data[0]['id']
                        supabase.table('estadisticas_usuarios').update({
                            'examenes_rendidos_hoy': stats_response.data[0].get('examenes_rendidos_hoy', 0) + 1,
                            'preguntas_correctas_hoy': stats_response.data[0].get('preguntas_correctas_hoy', 0) + correctas,
                            'preguntas_incorrectas_hoy': stats_response.data[0].get('preguntas_incorrectas_hoy', 0) + incorrectas,
                            'tiempo_total_estudio_hoy': stats_response.data[0].get('tiempo_total_estudio_hoy', 0) + int(float(resumen["tiempo_total"])),
                            'materias_estudiadas_hoy': [preguntas[0].get("tema", "General")]
                        }).eq('id', stats_id).execute()
                    else:
                        # Crear nuevas estadísticas para hoy
                        supabase.table('estadisticas_usuarios').insert({
                            'usuario_id': current_user['id'],
                            'fecha_estadistica': fecha_hoy,
                            'examenes_rendidos_hoy': 1,
                            'preguntas_correctas_hoy': correctas,
                            'preguntas_incorrectas_hoy': incorrectas,
                            'tiempo_total_estudio_hoy': int(float(resumen["tiempo_total"])),
                            'materias_estudiadas_hoy': [preguntas[0].get("tema", "General")]
                        }).execute()
                    
                    print(f"✅ Estadísticas diarias guardadas para {fecha_hoy}")
                except Exception as e:
                    print(f"⚠️ Error guardando estadísticas diarias: {e}")
                
                print(f"✅ Examen guardado exitosamente en Supabase para usuario {current_user['email']}")
                
        except Exception as e:
            print(f"❌ Error al guardar examen en Supabase: {e}")
            # Continuar sin guardar

    # OPCIONAL: seguir guardando en JSON para legacy
    with open("resultados.json", "a") as f:
        f.write(json.dumps(resumen) + "\n")

    return render_template("resultado_abierto.html", respuestas=respuestas, preguntas=preguntas, feedbacks=feedbacks, resumen=resumen, respuestas_texto_usuario=respuestas_texto_usuario, respuestas_texto_correcta=respuestas_texto_correcta)

@app.route("/cuestionario")
def cuestionario():
    # Verificar que hay preguntas en la sesión
    preguntas = session.get("preguntas", [])
    if not preguntas:
        flash("No hay un examen disponible para repetir. Genera un nuevo examen primero.", "error")
        return redirect(url_for("generar"))
    
    # Reiniciar respuestas y tiempos pero mantener las mismas preguntas
    session["respuestas"] = ["" for _ in preguntas]
    session["pregunta_times"] = []
    session["start_time"] = time.time()
    session["last_question_time"] = time.time()
    
    return redirect(url_for("pregunta", numero=0))

@app.route("/reiniciar", methods=["POST"])
def reiniciar():
    preguntas = session.get("preguntas", [])
    session["respuestas"] = ["" for _ in preguntas]
    session["pregunta_times"] = []
    session["start_time"] = time.time()
    session["last_question_time"] = time.time()
    return redirect(url_for("pregunta", numero=0))

@app.route("/historial")
def historial():
    # Verificar autenticación simple
    if not is_authenticated():
        return redirect(url_for('login'))
    
    try:
        if supabase:
            current_user = get_current_user()
            if not current_user:
                return redirect(url_for('login'))
            
            # Obtener exámenes del usuario
            examenes_response = supabase.table('examenes').select('*, carpetas(id, nombre, color)').eq('usuario_id', current_user['id']).order('fecha_rendido', desc=True).execute()
            
            # Obtener carpetas del usuario
            carpetas_response = supabase.table('carpetas').select('id, nombre, color').eq('usuario_id', current_user['id']).order('nombre').execute()
            
            # Calcular métricas del dashboard
            total_examenes = len(examenes_response.data) if examenes_response.data else 0
            total_carpetas = len(carpetas_response.data) if carpetas_response.data else 0
            
            # Calcular promedio de notas
            promedio_nota = 0
            tiempo_total_segundos = 0
            if examenes_response.data:
                notas = [examen.get('nota', 0) for examen in examenes_response.data if examen.get('nota') is not None]
                promedio_nota = sum(notas) / len(notas) if notas else 0
                # Filtrar valores None para el tiempo total
                tiempos = [examen.get('tiempo_duracion', 0) for examen in examenes_response.data if examen.get('tiempo_duracion') is not None]
                tiempo_total_segundos = sum(tiempos) if tiempos else 0
            
            # Convertir tiempo total a horas
            tiempo_total_horas = round(tiempo_total_segundos / 3600, 1)
            
            # Preparar exámenes recientes para el dashboard
            examenes_recientes = []
            if examenes_response.data:
                for examen in examenes_response.data[:5]:  # Solo los 5 más recientes
                    try:
                        # Formatear fecha para mostrar en hora de Argentina (GMT-3)
                        from datetime import timezone, timedelta
                        fecha_rendido = examen.get('fecha_rendido')
                        if fecha_rendido:
                            fecha_utc = datetime.fromisoformat(fecha_rendido.replace('Z', '+00:00'))
                            zona_horaria_argentina = timezone(timedelta(hours=-3))
                            fecha_argentina = fecha_utc.astimezone(zona_horaria_argentina)
                        else:
                            fecha_argentina = datetime.now()
                        
                        examenes_recientes.append({
                            'id': examen.get('id', 0),
                            'titulo': examen.get('titulo', 'Examen de General'),
                            'fecha': fecha_argentina,
                            'nota': examen.get('nota', 0),
                            'carpeta': examen.get('carpetas') if examen.get('carpetas') else None
                        })
                    except Exception as e:
                        continue
            
            # Preparar carpetas para el dashboard
            carpetas = []
            if carpetas_response.data:
                for carpeta in carpetas_response.data:
                    try:
                        # Contar exámenes en cada carpeta
                        examenes_carpeta_response = supabase.table('examenes').select('id', count='exact').eq('carpeta_id', carpeta.get('id')).execute()
                        cantidad_examenes = examenes_carpeta_response.count if hasattr(examenes_carpeta_response, 'count') else 0
                        
                        carpetas.append({
                            'id': carpeta.get('id', 0),
                            'nombre': carpeta.get('nombre', 'Sin nombre'),
                            'color': carpeta.get('color', '#6366f1'),
                            'cantidad_examenes': cantidad_examenes
                        })
                    except Exception as e:
                        continue
            
            # Renderizar el dashboard
            return render_template("historial.html", 
                total_examenes=total_examenes,
                promedio_nota=promedio_nota,
                tiempo_total_horas=tiempo_total_horas,
                total_carpetas=total_carpetas,
                examenes_recientes=examenes_recientes,
                carpetas=carpetas
            )
                
    except Exception as e:
        flash("Error al cargar el dashboard. Intenta nuevamente.")
        # En caso de error, mostrar dashboard con valores por defecto
        return render_template("historial.html", 
            total_examenes=0,
            promedio_nota=0,
            tiempo_total_horas=0,
            total_carpetas=0,
            examenes_recientes=[],
            carpetas=[]
        )

@app.route("/historial_completo")
def historial_completo():
    # Verificar autenticación simple
    if not is_authenticated():
        return redirect(url_for('login'))
    try:
        if supabase:
            # Obtener exámenes del usuario desde Supabase
            current_user = get_current_user()
            
            response = supabase.table('examenes').select('*, carpetas(id, nombre, color)').eq('usuario_id', current_user['id']).order('fecha_rendido', desc=True).execute()
            
            if response.data:
                examenes = []
                for examen in response.data:
                    # Formatear fecha para mostrar en hora de Argentina (GMT-3)
                    from datetime import timezone, timedelta
                    fecha_utc = datetime.fromisoformat(examen['fecha_rendido'].replace('Z', '+00:00'))
                    zona_horaria_argentina = timezone(timedelta(hours=-3))
                    fecha_argentina = fecha_utc.astimezone(zona_horaria_argentina)
                    
                    examenes.append({
                        'id': examen['id'],
                        'titulo': examen.get('titulo', 'Examen de General'),
                        'fecha': fecha_argentina,
                        'nota': examen['nota'],
                        'materia': examen['materia'],
                        'tiempo_total': examen['tiempo_duracion'],
                        'estado': examen['estado'],
                        # Agregar métricas detalladas
                        'correctas': examen.get('correctas', 0),
                        'parciales': examen.get('parciales', 0),
                        'incorrectas': examen.get('incorrectas', 0),
                        # Agregar información de carpeta
                        'carpeta': examen.get('carpetas') if examen.get('carpetas') else None
                    })
                return render_template("historial_completo.html", examenes=examenes)
            else:
                return render_template("historial_completo.html", examenes=[])
                
    except Exception as e:
        flash("Error al cargar el historial completo. Intenta nuevamente.")
        return redirect(url_for('generar'))

@app.route("/examen/<examen_id>")
def detalle_examen(examen_id):
    # Verificar autenticación simple
    if not is_authenticated():
        return redirect(url_for('login'))
    try:
        if supabase:
            # Obtener examen desde Supabase
            current_user = get_current_user()
            examen_response = supabase.table('examenes').select('*').eq('id', examen_id).eq('usuario_id', current_user['id']).execute()
            
            if examen_response.data:
                examen = examen_response.data[0]
                
                # Obtener preguntas del examen
                preguntas_response = supabase.table('preguntas_examen').select('*').eq('examen_id', examen_id).order('orden').execute()
                
                if preguntas_response.data:
                    preguntas = []
                    for pregunta in preguntas_response.data:
                        # Decodificar opciones JSON
                        opciones = []
                        if pregunta.get('opciones'):
                            try:
                                opciones = json.loads(pregunta['opciones'])
                            except:
                                opciones = []
                        
                        # Determinar el feedback basado en la respuesta
                        feedback = ""
                        if pregunta.get('respuesta_usuario') == pregunta.get('respuesta_correcta'):
                            feedback = "✔️ CORRECTA"
                        else:
                            feedback = "❌ INCORRECTA"
                        
                        preguntas.append({
                            'enunciado': pregunta['pregunta'],
                            'opciones': opciones,
                            'opciones_decoded': opciones,  # Para el template
                            'respuesta_usuario': pregunta['respuesta_usuario'],
                            'respuesta_correcta': pregunta['respuesta_correcta'],
                            'tipo': 'multiple_choice',  # Por defecto
                            'tema': 'General',  # Por defecto
                            'feedback': feedback
                        })
                    
                    # Formatear examen para el template
                    from datetime import timezone, timedelta
                    
                    # Convertir UTC a hora de Argentina (GMT-3)
                    fecha_utc = datetime.fromisoformat(examen['fecha_rendido'].replace('Z', '+00:00'))
                    zona_horaria_argentina = timezone(timedelta(hours=-3))
                    fecha_argentina = fecha_utc.astimezone(zona_horaria_argentina)
                    
                    examen_formateado = {
                        'id': examen['id'],
                        'titulo': examen.get('titulo', 'Examen de General'),
                        'fecha': fecha_argentina,
                        'nota': examen['nota'],
                        'materia': examen['materia'],
                        'tiempo_total': examen['tiempo_duracion'],
                        # Agregar métricas detalladas
                        'correctas': examen.get('correctas', 0),
                        'parciales': examen.get('parciales', 0),
                        'incorrectas': examen.get('incorrectas', 0),
                        'feedback_general': examen.get('feedback_general', 'Sin feedback disponible')
                    }
                    
                    return render_template("detalle_examen.html", examen=examen_formateado, preguntas=preguntas)
                else:
                    flash("No se encontraron preguntas para este examen.")
                    return redirect(url_for('historial'))
            else:
                flash("Examen no encontrado o no tienes acceso.")
                return redirect(url_for('historial'))
                
    except Exception as e:
        print(f"Error obteniendo detalle del examen: {e}")
        flash("Error al cargar el examen. Intenta nuevamente.")
        return redirect(url_for('historial'))

@app.route("/wolfram", methods=["GET", "POST"])
def wolfram_query():
    print(f"[DEBUG] Wolfram query - Method: {request.method}")
    # Verificar autenticación simple
    if not is_authenticated():
        print("[DEBUG] Usuario no autenticado")
        return redirect(url_for('login'))
    resultado = None
    imagen_url = None
    error = None
    pods = []
    if request.method == "POST":
        print(f"[DEBUG] Form data: {request.form}")
        print(f"[DEBUG] Operación: {request.form.get('operacion', '')}")
        print(f"[DEBUG] Expresión: {request.form.get('expresion', '')}")
        operacion = request.form.get("operacion", "")
        expresion = request.form.get("expresion", "")
        consulta = expresion.strip()
        # Si el usuario eligió una operación, armar la consulta
        if operacion and operacion != "":
            if operacion == "derivative":
                consulta = f"derivative of {expresion}"
            elif operacion == "integral":
                consulta = f"integrate {expresion}"
            elif operacion == "solve":
                consulta = f"solve {expresion}"
            elif operacion == "limit":
                consulta = f"limit {expresion}"
            elif operacion == "simplify":
                consulta = f"simplify {expresion}"
            elif operacion == "expand":
                consulta = f"expand {expresion}"
            elif operacion == "factor":
                consulta = f"factor {expresion}"
            elif operacion == "plot":
                consulta = f"plot {expresion}"
        # Si la consulta parece una frase, traducir con IA
        elif len(expresion.split()) > 4:
            try:
                prompt_ia = (
                    "Convertí la siguiente frase a una consulta matemática en inglés para Wolfram Alpha. "
                    "No expliques, solo devolvé la consulta lista para enviar.\nFrase: " + expresion
                )
                print(f"🔍 [WOLFRAM] Traduciendo expresión matemática...")
                print(f"🔍 [WOLFRAM] Longitud del prompt: {len(prompt_ia)} caracteres")
                
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": "Sos un traductor de frases matemáticas a consultas para Wolfram Alpha."},
                        {"role": "user", "content": prompt_ia}
                    ],
                    max_completion_tokens=60,
                    timeout=30
                )
                
                # Log de tokens utilizados
                if hasattr(response, 'usage'):
                    print(f"🔍 [WOLFRAM] Tokens utilizados:")
                    print(f"   📥 Input tokens: {response.usage.prompt_tokens}")
                    print(f"   📤 Output tokens: {response.usage.completion_tokens}")
                    print(f"   📊 Total tokens: {response.usage.total_tokens}")
                
                consulta = response.choices[0].message.content.strip()
                print(f"🔍 [WOLFRAM] Consulta generada: {len(consulta)} caracteres")
            except Exception as e:
                error = f"No se pudo traducir la frase a consulta matemática: {str(e)}"
        try:
            print(f"[DEBUG] Consulta a Wolfram: {consulta}")
            print(f"[DEBUG] API Key: {app_id}")
            
            url = "https://api.wolframalpha.com/v2/query"
            params = {
                "input": consulta,
                "appid": app_id,
                "format": "image,plaintext"
            }
            print(f"[DEBUG] Parámetros: {params}")
            
            resp = requests.get(url, params=params, timeout=30)
            print(f"[DEBUG] Status Code: {resp.status_code}")
            print(f"[DEBUG] Response: {resp.text[:500]}...")
            
            if resp.status_code != 200:
                error = f"Error HTTP: {resp.status_code}"
                print(f"[DEBUG] Error HTTP: {resp.status_code}")
            else:
                root = ET.fromstring(resp.text)
                for pod in root.findall(".//pod"):
                    pod_title = pod.attrib.get("title", "")
                    subpod = pod.find("subpod")
                    pod_plaintext = subpod.findtext("plaintext") if subpod is not None else None
                    pod_img = None
                    if subpod is not None:
                        img_tag = subpod.find("img")
                        if img_tag is not None:
                            pod_img = img_tag.attrib.get("src")
                    # Guardar todos los pods relevantes
                    if pod_plaintext or pod_img:
                        pods.append({"title": pod_title, "plaintext": pod_plaintext, "img": pod_img})
                        print(f"[DEBUG] Pod encontrado: {pod_title} - {pod_plaintext}")
                    # Guardar el resultado principal
                    if pod_title.lower() in ["result", "resultado", "solution", "solución"] and not resultado:
                        resultado = pod_plaintext
                        imagen_url = pod_img
                        print(f"[DEBUG] Resultado principal: {resultado}")
                if not resultado and pods:
                    resultado = pods[0]["plaintext"]
                    imagen_url = pods[0]["img"]
                    print(f"[DEBUG] Usando primer pod como resultado: {resultado}")
                if not resultado:
                    error = "No se encontró una respuesta clara para tu consulta."
                    print("[DEBUG] No se encontró resultado")
        except Exception as e:
            import traceback
            print("Error Wolfram:", e)
            traceback.print_exc()
            error = f"Error al consultar Wolfram Alpha: {str(e)}"
    
    print(f"[DEBUG] Final - Resultado: {resultado}")
    print(f"[DEBUG] Final - Error: {error}")
    print(f"[DEBUG] Final - Pods: {len(pods)}")
    
    return render_template("wolfram.html", resultado=resultado, imagen_url=imagen_url, error=error, pods=pods)

@app.route("/examen_matematico/<int:numero>", methods=["GET", "POST"])
def examen_matematico(numero):
    # Verificar autenticación simple
    if not is_authenticated():
        return redirect(url_for('login'))
    ejercicios = session.get("ejercicios_matematicos", [])
    if not ejercicios or numero >= len(ejercicios):
        return redirect(url_for("resultado_matematico"))

    if request.method == "POST":
        respuesta = request.form.get("respuesta", "")
        ejercicios[numero]["respuesta_usuario"] = respuesta
        session["ejercicios_matematicos"] = ejercicios
        return redirect(url_for("examen_matematico", numero=numero + 1))

    ejercicio = ejercicios[numero]
    return render_template("examen_matematico.html", ejercicio=ejercicio, numero=numero + 1, total=len(ejercicios), actual=numero)

@app.route("/resultado_matematico")
def resultado_matematico():
    # Verificar autenticación simple
    if not is_authenticated():
        return redirect(url_for('login'))
    ejercicios = session.get("ejercicios_matematicos", [])
    for ejercicio in ejercicios:
        usuario = ejercicio.get("respuesta_usuario", "").strip()
        solucion = ejercicio.get("solucion", "").strip()
        es_correcta = False
        if usuario and solucion:
            try:
                url = "https://api.wolframalpha.com/v2/query"
                # Normalizar respuestas
                usuario_norm = usuario.replace(" ", "").lower()
                solucion_norm = solucion.replace(" ", "").lower()
                # 1. Consulta directa
                consulta_equiv = f"is ({usuario}) = ({solucion})"
                params = {"input": consulta_equiv, "appid": app_id, "format": "plaintext"}
                resp = requests.get(url, params=params, timeout=30)
                if resp.status_code == 200:
                    import xml.etree.ElementTree as ET
                    root = ET.fromstring(resp.text)
                    for pod in root.findall(".//pod"):
                        pod_title = pod.attrib.get("title", "").lower()
                        if pod_title in ["result", "resultado", "solution", "solución"]:
                            subpod = pod.find("subpod")
                            result_text = subpod.findtext("plaintext") if subpod is not None else ""
                            if result_text and "true" in result_text.lower():
                                es_correcta = True
                # 2. Solo parte derecha de la ecuación (si hay '=')
                if not es_correcta and '=' in solucion:
                    derecha = solucion.split('=')[-1].strip()
                    consulta_equiv2 = f"is ({usuario}) = ({derecha})"
                    params2 = {"input": consulta_equiv2, "appid": app_id, "format": "plaintext"}
                    resp2 = requests.get(url, params=params2, timeout=30)
                    if resp2.status_code == 200:
                        root2 = ET.fromstring(resp2.text)
                        for pod in root2.findall(".//pod"):
                            pod_title = pod.attrib.get("title", "").lower()
                            if pod_title in ["result", "resultado", "solution", "solución"]:
                                subpod = pod.find("subpod")
                                result_text = subpod.findtext("plaintext") if subpod is not None else ""
                                if result_text and "true" in result_text.lower():
                                    es_correcta = True
                # 3. Normalizar y comparar texto plano (fallback)
                if not es_correcta and usuario_norm == solucion_norm:
                    es_correcta = True
                ejercicio["es_correcta"] = es_correcta
                print(f"[WOLFRAM] Ejercicio: {ejercicio.get('enunciado','')}")
                print(f"[WOLFRAM] Resultado correcto (texto plano): {solucion}")
            except Exception as e:
                ejercicio["es_correcta"] = False
        else:
            ejercicio["es_correcta"] = False
    return render_template("resultado_matematico.html", ejercicios=ejercicios)

@app.route('/como-funciona')
def como_funciona():
    return render_template('como_funciona.html')

@app.route("/planificacion", methods=["GET", "POST"])
def planificacion():
    if request.method == "POST":
        fecha_examen = request.form.get("fecha_examen")
        dias_no = request.form.get("dias_no", "")
        dias_no_multiple = request.form.get("dias_no_multiple", "")
        tiempo_dia = request.form.get("tiempo_dia")
        aclaraciones = request.form.get("aclaraciones", "")
        resumen = request.form.get("resumen", "")
        archivo = request.files.get("archivo")
        texto_resumen = resumen.strip()
        # Procesar archivo si existe (copiado exactamente del generador de exámenes)
        if archivo and archivo.filename:
            if archivo.filename.endswith(".txt"):
                texto_resumen = archivo.read().decode("utf-8")
            elif archivo.filename.endswith(".pdf"):
                from io import BytesIO
                import PyPDF2
                pdf_stream = BytesIO(archivo.read())
                reader = PyPDF2.PdfReader(pdf_stream)
                texto_resumen = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
            elif archivo.filename.endswith(".docx"):
                from io import BytesIO
                import docx
                docx_stream = BytesIO(archivo.read())
                doc = docx.Document(docx_stream)
                texto_resumen = "\n".join([p.text for p in doc.paragraphs])
        print("\n--- TEXTO EXTRAÍDO DEL ARCHIVO (PLANIFICACIÓN) ---\n", texto_resumen, "\n--- FIN TEXTO EXTRAÍDO ---\n")
        # Procesar días no disponibles (usar múltiples fechas si están disponibles)
        print(f"🔍 DEBUG - dias_no: '{dias_no}'")
        print(f"🔍 DEBUG - dias_no_multiple: '{dias_no_multiple}'")
        
        dias_no_final = dias_no
        if dias_no_multiple:
            dias_no_final = dias_no_multiple
        elif dias_no:
            dias_no_final = dias_no
            

        
        # Preparar datos para el nuevo prompt optimizado
        from datetime import date, timedelta
        fecha_actual = date.today().strftime('%Y-%m-%d')
        
        # Procesar días no disponibles para el nuevo formato
        dias_no_disponibles_formateados = []
        if dias_no_final:
            # Separar por comas y limpiar
            dias_lista = [dia.strip() for dia in dias_no_final.split(',') if dia.strip()]
            for dia in dias_lista:
                # Si es una fecha (contiene -), mantenerla
                if '-' in dia:
                    dias_no_disponibles_formateados.append(dia)
                # Si es un día de la semana, mantenerlo
                else:
                    dias_no_disponibles_formateados.append(dia.lower())
        
        # Preparar horas por día (convertir a número si es string)
        try:
            horas_por_dia = float(tiempo_dia) if tiempo_dia else 3.0
        except:
            horas_por_dia = 3.0
        
        # Nuevo prompt optimizado con GPT-5.0
        prompt_template = """Rol e intención

Eres un planificador de estudios estricto y fiable. Tu única salida debe ser un JSON válido que describa un plan de estudio diario que respeta absolutamente todas las restricciones del usuario.
Es fundamental que TODO el temario quede cubierto, sin omitir ningún tema o subtema. Si el tiempo no alcanza para mantenerlos separados, agrúpalos de manera explícita, pero jamás los elimines.

Entradas (proporcionadas por la aplicación)

temario_texto: texto plano completo extraído del archivo del usuario (PDF/DOCX). Mantiene saltos de línea y viñetas.
fecha_inicio: fecha ISO YYYY-MM-DD desde la que se puede empezar a estudiar.
fecha_examen: fecha ISO YYYY-MM-DD del examen (el plan debe terminar ese día o antes).
dias_no_disponibles: lista que puede incluir fechas ISO específicas [\"YYYY-MM-DD\", ...] y/o nombres de días de la semana en español [\"lunes\", \"martes\", ...].
horas_por_dia: puede ser
un número (horas fijas por cada día disponible), o
un objeto {{ "YYYY-MM-DD": horas, ... }} para horas variables por fecha.
aclaraciones_adicionales: texto libre con prioridades, temas clave, métodos preferidos, etc. Úsalo para ponderar la importancia/orden.

Objetivo

Generar un plan detallado, específico y realista que:
- Termine en o antes de fecha_examen.
- Excluya por completo los dias_no_disponibles (fechas y/o días de la semana).
- Distribuya TODO el temario, sin dejar nada afuera. Si es necesario, agrupa subtemas relacionados en un solo bloque.
- Respete estrictamente horas_por_dia, permitiendo dividir la carga de un mismo día en varias actividades.
- Incluya TODOS los temas del temario (sin omitir ninguno).
- Distinga claramente entre temas principales y subtemas.
- Use exactamente el formato de actividad: "Estudiar [Tema Principal] | [Subtema 1], [Subtema 2], ...".
- Añada días de repaso antes del examen.
- Añada mensajes motivacionales solo en días con carga alta sin romper el formato exigido.

Detección y estructuración de contenidos del temario

- Analiza temario_texto y extrae TODOS los temas, conservando su orden.
- Identifica temas principales usando numeraciones, mayúsculas, encabezados.
- Extrae subtemas de ítems subordinados.
- Si hay más de dos niveles, mapea: Nivel 1 → tema_principal; Niveles 2+ → subtemas.
- No inventes temas. Incluye todo lo que aparezca en el temario.

Asignación temporal y realismo

- Genera un calendario entre fecha_inicio y fecha_examen excluyendo los días prohibidos.
- Capacidad diaria: calcula subtemas por día según horas_por_dia, usando la regla 1 subtema ≈ 45 minutos.
- Si un tema principal tiene demasiados subtemas para un solo día, divídelo en varios días consecutivos, pero no lo recortes.
- Si el total de subtemas excede el tiempo disponible, AGRUPA, nunca omitas. Usa nombres compuestos claros (ejemplo: "2.2–2.4 Regresiones: lineal, logística, regularización").
- Está permitido generar más de una entrada de actividad por día, mientras el total no exceda la capacidad diaria.

Repaso

- Incluye 1–2 días de repaso antes del examen.
- Usa como tema_principal “Repaso general” o repaso por bloques.
- Subtemas: lista de los puntos clave o difíciles.

Mensajes motivacionales

- Solo si carga ≥ 4h o ≥ 6 subtemas.
- Añádelos al final de “actividad” tras “ // ”.
- Breves y en español.

Reglas duras

- No programes nada fuera del rango [fecha_inicio, fecha_examen].
- No incluyas entradas en dias_no_disponibles.
- No excedas capacidad diaria. Divide en varias actividades si es necesario.
- Incluye TODO el contenido del temario. Si falta tiempo, agrupa explícitamente.
- Mantén el orden lógico del temario. Solo reordena mínimamente si es imprescindible para equilibrar carga.
- Formato EXACTO de actividad: "Estudiar [Tema Principal] | [Subtema 1], [Subtema 2], ..." (+ opcional " // [Mensaje]").
- El último día del plan debe ser en o antes de fecha_examen.
- Salida estrictamente en español.

Formato de salida (JSON estricto)

Devuelve únicamente un array JSON. Sin texto adicional antes o después. Sin comentarios. Sin comas finales. Sin claves extra.
Cada elemento del array debe tener exactamente estas claves:
"fecha": "YYYY-MM-DD"
"actividad": "Estudiar [Tema Principal] | [Subtema 1], [Subtema 2], ..." + opcional " // [Mensaje motivacional]"
"tema_principal": "[Tema Principal]"
"subtemas": ["[Subtema 1]", "[Subtema 2]", ...]
Ordena el array ascendentemente por "fecha".
Permite múltiples elementos con la misma fecha, siempre que no superen en total las horas disponibles de ese día.

Datos
temario_texto: {temario_texto}
fecha_inicio: {fecha_inicio}
fecha_examen: {fecha_examen}
dias_no_disponibles: {dias_no_disponibles}
horas_por_dia: {horas_por_dia}
aclaraciones_adicionales: {aclaraciones_adicionales}

"""
        
        # Consultar a OpenAI con GPT-4.1 y response_format
        try:
            # Construir el prompt con las variables
            prompt = prompt_template.format(
                temario_texto=texto_resumen,
                fecha_inicio=fecha_actual,
                fecha_examen=fecha_examen,
                dias_no_disponibles=dias_no_disponibles_formateados,
                horas_por_dia=horas_por_dia,
                aclaraciones_adicionales=aclaraciones
            )
            print(f"🔍 [PLANIFICACIÓN] Enviando prompt a GPT-4o...")
            print(f"🔍 [PLANIFICACIÓN] Longitud del prompt: {len(prompt)} caracteres")
            print(f"🔍 [PLANIFICACIÓN] Días no disponibles: {dias_no_disponibles_formateados}")
            print(f"🔍 [PLANIFICACIÓN] Horas por día: {horas_por_dia}")
            print(f"🔍 [PLANIFICACIÓN] Template antes del format:")
            print(f"🔍 [PLANIFICACIÓN] {prompt_template[:500]}...")
            print(f"🔍 [PLANIFICACIÓN] Variables del format:")
            print(f"🔍 [PLANIFICACIÓN] temario_texto: {len(texto_resumen)} chars")
            print(f"🔍 [PLANIFICACIÓN] fecha_inicio: {fecha_actual}")
            print(f"🔍 [PLANIFICACIÓN] fecha_examen: {fecha_examen}")
            print(f"🔍 [PLANIFICACIÓN] dias_no_disponibles: {dias_no_disponibles_formateados}")
            print(f"🔍 [PLANIFICACIÓN] horas_por_dia: {horas_por_dia}")
            print(f"🔍 [PLANIFICACIÓN] aclaraciones_adicionales: {aclaraciones}")
            
            # Log del prompt completo antes de enviarlo
            print(f"🔍 [PLANIFICACIÓN] PROMPT COMPLETO:")
            print(f"🔍 [PLANIFICACIÓN] {prompt}")
            print(f"🔍 [PLANIFICACIÓN] --- FIN PROMPT ---")
            
            response = client.chat.completions.create(
                model="gpt-4o", 
                messages=[
                    {"role": "user", "content": prompt}
                ],
                max_completion_tokens=6000,
                temperature=0.2,  # Para consistencia
                seed=123,  # Para reproducibilidad
                response_format={
                    "type": "json_schema",
                    "json_schema": {
                        "name": "plan_estudio",
                        "schema": {
                            "type": "object",
                            "additionalProperties": False,
                            "required": ["plan"],
                            "properties": {
                                "plan": {
                                    "type": "array",
                                    "items": {
                                        "type": "object",
                                        "additionalProperties": False,
                                        "required": ["fecha", "actividad", "tema_principal", "subtemas"],
                                        "properties": {
                                            "fecha": {"type": "string", "pattern": "^\\d{4}-\\d{2}-\\d{2}$"},
                                            "actividad": {"type": "string"},
                                            "tema_principal": {"type": "string"},
                                            "subtemas": {
                                                "type": "array",
                                                "items": {"type": "string"},
                                                "minItems": 1
                                            }
                                        }
                                    }
                                }
                            }
                        },
                        "strict": True
                    }
                },
                timeout=120
            )
            
            # Log de tokens utilizados
            if hasattr(response, 'usage'):
                print(f"🔍 [PLANIFICACIÓN] Tokens utilizados:")
                print(f"   📥 Input tokens: {response.usage.prompt_tokens}")
                print(f"   📤 Output tokens: {response.usage.completion_tokens}")
                print(f"   📊 Total tokens: {response.usage.total_tokens}")
            
            # Debug: verificar si tiene parsed
            print(f"🔍 [PLANIFICACIÓN] hasattr parsed: {hasattr(response.choices[0].message, 'parsed')}")
            print(f"🔍 [PLANIFICACIÓN] message attributes: {dir(response.choices[0].message)}")
            
            # Obtener el JSON parseado directamente
            if hasattr(response.choices[0].message, 'parsed'):
                parsed_data = response.choices[0].message.parsed
                print(f"🔍 [PLANIFICACIÓN] parsed_data: {type(parsed_data)}")
                print(f"🔍 [PLANIFICACIÓN] parsed_data keys: {list(parsed_data.keys()) if isinstance(parsed_data, dict) else 'No es dict'}")
                print(f"🔍 [PLANIFICACIÓN] parsed_data content: {parsed_data}")
                
                # Extraer el array del objeto
                if isinstance(parsed_data, dict) and 'plan' in parsed_data:
                    plan = parsed_data['plan']
                elif isinstance(parsed_data, list):
                    plan = parsed_data
                else:
                    plan = []
                
                plan_json = json.dumps(plan, ensure_ascii=False, indent=2)
                print(f"✅ [PLANIFICACIÓN] JSON parseado directamente por GPT-4o")
                print(f"🔍 [PLANIFICACIÓN] Plan obtenido: {len(plan)} actividades")
                print(f"🔍 [PLANIFICACIÓN] Plan type: {type(plan)}")
            else:
                # Fallback al método tradicional
                plan_json = response.choices[0].message.content.strip()
                print(f"✅ [PLANIFICACIÓN] Respuesta recibida de GPT-4o")
                print(f"🔍 [PLANIFICACIÓN] Contenido obtenido: {len(plan_json)} caracteres")
                
                # Log detallado de la respuesta
                print(f"🔍 [PLANIFICACIÓN] RESPUESTA COMPLETA:")
                print(f"🔍 [PLANIFICACIÓN] {plan_json}")
                print(f"🔍 [PLANIFICACIÓN] --- FIN RESPUESTA ---")
                
                if len(plan_json) == 0:
                    print("❌ [PLANIFICACIÓN] El contenido está vacío!")
                    plan_json = None
                else:
                    print(f"🔍 [PLANIFICACIÓN] Primeros 500 caracteres:")
                    print(f"   {plan_json[:500]}...")
        except Exception as e:
            plan_json = f"[{{'fecha':'error','actividad':'Error al generar planificación: {str(e)}'}}]"
        # Extraer temas/unidades del resumen (líneas no vacías con más de 10 caracteres)
        import json, re
        from datetime import datetime, timedelta, date
        explicacion_ia = None
        
        # Si ya tenemos el plan parseado desde response_format, usarlo directamente
        if 'plan' in locals() and plan is not None:
            print(f"✅ [PLANIFICACIÓN] Usando plan parseado directamente: {len(plan)} actividades")
        else:
            # Intentar parsear directamente el JSON
            try:
                parsed_json = json.loads(plan_json)
                print(f"✅ [PLANIFICACIÓN] JSON parseado correctamente")
                print(f"🔍 [PLANIFICACIÓN] Tipo de parsed_json: {type(parsed_json)}")
                
                # Extraer el array del objeto
                if isinstance(parsed_json, dict) and 'plan' in parsed_json:
                    plan = parsed_json['plan']
                    print(f"✅ [PLANIFICACIÓN] Array extraído del objeto: {len(plan)} actividades")
                elif isinstance(parsed_json, list):
                    plan = parsed_json
                    print(f"✅ [PLANIFICACIÓN] JSON es array directo: {len(plan)} actividades")
                else:
                    print(f"❌ [PLANIFICACIÓN] Estructura JSON no reconocida: {type(parsed_json)}")
                    plan = None
                    
                explicacion_ia = None  # Si se puede parsear como JSON, no hay explicación
            except json.JSONDecodeError as e:
                print(f"⚠️ [PLANIFICACIÓN] Error parseando JSON: {e}")
                print(f"⚠️ [PLANIFICACIÓN] Contenido que falló: {plan_json}")
                print(f"⚠️ [PLANIFICACIÓN] Intentando extraer...")
                # Si no es JSON válido, intentar extraer JSON con regex
                # Primero intentar extraer JSON de markdown (```json ... ```)
                markdown_match = re.search(r'```json\s*([\s\S]*?)\s*```', plan_json)
                if markdown_match:
                    json_str = markdown_match.group(1)
                    print(f"🔍 [PLANIFICACIÓN] JSON extraído de markdown: {json_str}")
                    try:
                        plan = json.loads(json_str)
                        explicacion_ia = None
                        print(f"✅ [PLANIFICACIÓN] JSON extraído de markdown: {len(plan)} actividades")
                    except Exception as e2:
                        plan = None
                        explicacion_ia = plan_json
                else:
                    # Intentar extraer JSON normal con regex
                    print(f"🔍 [PLANIFICACIÓN] Intentando extraer JSON con regex...")
                    match = re.search(r'\[\s*{[\s\S]*?}\s*\]', plan_json)
                    if match:
                        json_str = match.group(0)
                        print(f"🔍 [PLANIFICACIÓN] JSON extraído con regex: {json_str}")
                        try:
                            plan = json.loads(json_str)
                            print(f"✅ [PLANIFICACIÓN] JSON extraído con regex: {len(plan)} actividades")
                            # Si hay texto antes del JSON, lo guardo como explicación
                            if plan_json.strip() != json_str.strip():
                                explicacion_ia = plan_json.replace(json_str, '').strip()
                            else:
                                explicacion_ia = None
                        except Exception as e3:
                            print(f"❌ [PLANIFICACIÓN] Error parseando JSON extraído: {e3}")
                            plan = None
                            explicacion_ia = plan_json
                    else:
                        print(f"❌ [PLANIFICACIÓN] No se encontró JSON válido en la respuesta")
                        # Si no hay JSON, mostrar como texto plano
                        explicacion_ia = plan_json
                        plan = None
                        print(f"❌ [PLANIFICACIÓN] No se pudo parsear JSON, usando como explicación")
        
        # Verificar y procesar la estructura del plan
        if plan and len(plan) > 0:
            # Verificar si ya tiene la estructura nueva (tema_principal, subtemas)
            primer_item = plan[0]
            if 'tema_principal' in primer_item and 'subtemas' in primer_item:
                print(f"✅ [PLANIFICACIÓN] Plan ya tiene estructura completa (tema_principal, subtemas)")
            else:
                print(f"🔄 [PLANIFICACIÓN] Procesando plan para agregar tema_principal y subtemas")
                # Procesar cada actividad para separar tema principal y subtemas
                for item in plan:
                    actividad = item.get('actividad', '')
                    
                    # Extraer mensaje motivacional si existe (después de //)
                    if ' // ' in actividad:
                        actividad_base, mensaje_motivacional = actividad.split(' // ', 1)
                        item['actividad'] = actividad_base.strip()
                        item['mensaje_motivacional'] = mensaje_motivacional.strip()
                    else:
                        item['mensaje_motivacional'] = None
                    
                    # Procesar tema principal y subtemas
                    if '|' in actividad:
                        # Si hay múltiples temas separados por |, el primer tema es el principal
                        partes = actividad.split('|')
                        item['tema_principal'] = partes[0].strip().replace('Estudiar ', '').strip()
                        # Los temas restantes se convierten en subtemas
                        if len(partes) > 1:
                            item['subtemas'] = [parte.strip() for parte in partes[1:] if parte.strip()]
                        else:
                            item['subtemas'] = []
                    else:
                        # Si no hay |, toda la actividad es el tema principal
                        item['tema_principal'] = actividad.replace('Estudiar ', '').strip()
                        item['subtemas'] = []
        
        # Preparar datos para el calendario visual
        fechas_ordenadas = []
        actividades_por_fecha = {}
        
        if plan and isinstance(plan, list) and len(plan) > 0 and 'fecha' in plan[0]:
            # Crear diccionario de actividades por fecha
            actividades_por_fecha = {item['fecha']: item['actividad'] for item in plan}
            
            # Ordenar fechas y procesar para el template
            fechas_ordenadas = []
            for item in plan:
                try:
                    fecha_str = item.get('fecha', '')
                    actividad = item.get('actividad', '')
                    fecha_dt = datetime.strptime(fecha_str, '%Y-%m-%d')
                    fechas_ordenadas.append((fecha_dt, fecha_str, actividad))
                except ValueError:
                    continue
            
            fechas_ordenadas.sort(key=lambda x: x[0])
            print(f"✅ [PLANIFICACIÓN] Plan procesado: {len(fechas_ordenadas)} fechas válidas")

        

        
        return render_template("planificacion_resultado.html", plan=plan, plan_json=plan_json, actividades_por_fecha=actividades_por_fecha, explicacion_ia=explicacion_ia, fechas_ordenadas=fechas_ordenadas, es_planificacion_guardada=False)
    
    return render_template("planificacion.html")

@app.route("/guardar-planificacion", methods=["POST"])
def guardar_planificacion():
    """Guardar una planificación de estudio en la base de datos"""
    if not is_authenticated():
        return redirect(url_for('login'))
    
    try:
        current_user = get_current_user()
        
        # Obtener datos del formulario
        titulo = request.form.get("titulo", "Mi plan de estudio")
        fecha_examen = request.form.get("fecha_examen")
        dias_no = request.form.get("dias_no", "")
        dias_no_multiple = request.form.get("dias_no_multiple", "")
        tiempo_dia = request.form.get("tiempo_dia")
        aclaraciones = request.form.get("aclaraciones", "")
        plan_json = request.form.get("plan_json")
        explicacion_ia = request.form.get("explicacion_ia", "")
        
        # Validar datos requeridos
        if not all([fecha_examen, tiempo_dia, plan_json]):
            flash("Faltan datos requeridos para guardar la planificación")
            return redirect(url_for('planificacion'))
        
        # Procesar días no disponibles (usar múltiples fechas si están disponibles)
        print(f"🔍 GUARDAR DEBUG - dias_no: '{dias_no}'")
        print(f"🔍 GUARDAR DEBUG - dias_no_multiple: '{dias_no_multiple}'")
        
        dias_no_disponibles = []
        dias_no_final = dias_no
        if dias_no_multiple:
            dias_no_final = dias_no_multiple
        elif dias_no:
            dias_no_final = dias_no
            
        print(f"🔍 GUARDAR DEBUG - dias_no_final: '{dias_no_final}'")
            
        if dias_no_final:
            dias_no_disponibles = [dia.strip() for dia in dias_no_final.split(',') if dia.strip()]
            
        print(f"🔍 GUARDAR DEBUG - dias_no_disponibles: {dias_no_disponibles}")
        
        # Crear planificación en la base de datos
        nueva_planificacion = {
            'usuario_id': current_user['id'],
            'titulo': titulo,
            'fecha_examen': fecha_examen,
            'dias_no_disponibles': dias_no_disponibles,
            'tiempo_por_dia': float(tiempo_dia),
            'aclaraciones': aclaraciones,
            'plan_json': plan_json,
            'explicacion_ia': explicacion_ia,
            'ultima_actividad': datetime.utcnow().isoformat()
        }
        
        print(f"🔍 Intentando guardar planificación: {nueva_planificacion}")
        
        # Usar el usuario autenticado en Supabase para evitar problemas de RLS
        try:
            response = supabase.table('planificaciones').insert(nueva_planificacion).execute()
            print(f"✅ Respuesta de Supabase: {response}")
        except Exception as e:
            print(f"❌ Error en Supabase: {e}")
            # Intentar con el usuario autenticado
            response = supabase.auth.get_user()
            if response.user:
                print(f"🔍 Usuario autenticado en Supabase: {response.user.id}")
                response = supabase.table('planificaciones').insert(nueva_planificacion).execute()
            else:
                raise e
        
        if response.data:
            print(f"✅ Planificación guardada para usuario {current_user['email']}")
            flash("¡Planificación guardada exitosamente! Puedes verla en 'Mi Calendario'")
            return redirect(url_for('mi_calendario'))
        else:
            print(f"❌ Error guardando planificación para usuario {current_user['email']}")
            flash("Error al guardar la planificación. Intenta de nuevo.")
            return redirect(url_for('planificacion'))
            
    except Exception as e:
        print(f"❌ Error en guardar_planificacion: {e}")
        flash("Error interno al guardar la planificación")
        return redirect(url_for('planificacion'))

@app.route("/mi-calendario")
def mi_calendario():
    """Mostrar el calendario personal del usuario con todas sus planificaciones"""
    if not is_authenticated():
        return redirect(url_for('login'))
    
    try:
        current_user = get_current_user()
        
        # Obtener todas las planificaciones del usuario
        response = supabase.table('planificaciones').select('*').eq('usuario_id', current_user['id']).order('fecha_creacion', desc=True).execute()
        
        planificaciones = response.data if response.data else []
        
        print(f"📅 Planificaciones encontradas para usuario {current_user['email']}: {len(planificaciones)}")
        
        return render_template("mi_calendario.html", planificaciones=planificaciones)
        
    except Exception as e:
        print(f"❌ Error en mi_calendario: {e}")
        flash("Error al cargar tu calendario personal")
        return redirect(url_for('generar'))

@app.route("/ver-planificacion/<plan_id>")
def ver_planificacion(plan_id):
    """Mostrar una planificación específica guardada"""
    if not is_authenticated():
        return redirect(url_for('login'))
    
    try:
        current_user = get_current_user()
        
        # Obtener la planificación específica
        response = supabase.table('planificaciones').select('*').eq('id', plan_id).eq('usuario_id', current_user['id']).execute()
        
        if not response.data:
            flash("Planificación no encontrada")
            return redirect(url_for('mi_calendario'))
        
        planificacion = response.data[0]
        
        # Convertir el plan_json de vuelta a formato usable
        import json
        try:
            print(f"🔍 Plan JSON tipo: {type(planificacion['plan_json'])}")
            print(f"🔍 Plan JSON contenido: {planificacion['plan_json']}")
            
            # El plan_json puede venir en formato string o ya como objeto
            if isinstance(planificacion['plan_json'], str):
                # Limpiar el string de markdown si viene con ```json
                clean_json = planificacion['plan_json']
                if clean_json.startswith('```json'):
                    clean_json = clean_json.replace('```json', '').replace('```', '').strip()
                elif clean_json.startswith('```'):
                    clean_json = clean_json.replace('```', '').strip()
                
                print(f"🔍 JSON limpio: {clean_json}")
                plan_data = json.loads(clean_json)
            else:
                plan_data = planificacion['plan_json']
                
            print(f"🔍 Plan data después del parsing: {plan_data}")
            print(f"🔍 Plan data tipo: {type(plan_data)}")
            print(f"🔍 Plan data es lista: {isinstance(plan_data, list)}")
            if isinstance(plan_data, list):
                print(f"🔍 Plan data longitud: {len(plan_data)}")
                if len(plan_data) > 0:
                    print(f"🔍 Primer elemento: {plan_data[0]}")
                    print(f"🔍 Primer elemento tipo: {type(plan_data[0])}")
        except Exception as e:
            print(f"❌ Error parseando JSON: {e}")
            import traceback
            traceback.print_exc()
            plan_data = []
        
        print(f"🔍 Plan JSON original: {planificacion['plan_json']}")
        print(f"🔍 Plan data procesado: {plan_data}")
        
        # Preparar datos para el template usando la misma lógica del historial
        plan = []
        
        # El plan_data puede venir como diccionario con clave 'plan' o como lista directa
        plan_list = []
        if isinstance(plan_data, dict) and 'plan' in plan_data:
            plan_list = plan_data['plan']
            print(f"🔍 Extrayendo lista del diccionario: {len(plan_list)} elementos")
        elif isinstance(plan_data, list):
            plan_list = plan_data
            print(f"🔍 Usando lista directa: {len(plan_list)} elementos")
        else:
            print(f"🔍 Tipo de plan_data no reconocido: {type(plan_data)}")
        
        if plan_list and isinstance(plan_list, list):
            for item in plan_list:
                if isinstance(item, dict) and 'fecha' in item and 'actividad' in item:
                    # Extraer tema principal y subtemas
                    actividad = item['actividad']
                    if '|' in actividad:
                        partes = actividad.split('|')
                        tema_principal = partes[0].strip()
                        subtemas = [s.strip() for s in partes[1].split(',')] if len(partes) > 1 else []
                    else:
                        tema_principal = actividad
                        subtemas = []
                    
                    plan.append({
                        'fecha': item['fecha'],
                        'actividad': actividad,
                        'tema_principal': tema_principal,
                        'subtemas': subtemas
                    })
        
        print(f"📋 Plan procesado final: {len(plan)} elementos")
        for p in plan[:3]:  # Mostrar solo los primeros 3 para debug
            print(f"  - {p['fecha']}: {p['tema_principal']}")
        
        print(f"📋 Plan final para template: {plan}")
        print(f"📋 Mostrando planificación {plan_id} para usuario {current_user['email']}")
        
        return render_template("planificacion_resultado.html", 
                             plan=plan,
                             explicacion_ia=planificacion.get('explicacion_ia'),
                             plan_json=planificacion['plan_json'],
                             fecha_examen=planificacion['fecha_examen'],
                             dias_no=','.join(planificacion.get('dias_no_disponibles', [])),
                             tiempo_dia=planificacion['tiempo_por_dia'],
                             aclaraciones=planificacion.get('aclaraciones', ''),
                             es_planificacion_guardada=True,
                             days_list=[],
                             actividades_por_fecha={},
                             fechas_ordenadas=[])
        
    except Exception as e:
        print(f"❌ Error en ver_planificacion: {e}")
        flash("Error al cargar la planificación")
        return redirect(url_for('mi_calendario'))

@app.errorhandler(404)
def pagina_no_encontrada(e):
    return render_template('404.html'), 404

@app.route("/eliminar-planificacion/<plan_id>", methods=["POST"])
def eliminar_planificacion(plan_id):
    """Eliminar una planificación específica"""
    if not is_authenticated():
        return redirect(url_for('login'))
    
    try:
        current_user = get_current_user()
        
        # Verificar que la planificación pertenece al usuario
        response = supabase.table('planificaciones').select('id, titulo').eq('id', plan_id).eq('usuario_id', current_user['id']).execute()
        
        if not response.data:
            flash("Planificación no encontrada")
            return redirect(url_for('mi_calendario'))
        
        # Eliminar la planificación
        delete_response = supabase.table('planificaciones').delete().eq('id', plan_id).eq('usuario_id', current_user['id']).execute()
        
        if delete_response.data:
            flash("✅ Planificación eliminada exitosamente")
            print(f"🗑️ Planificación {plan_id} eliminada para usuario {current_user['email']}")
        else:
            flash("❌ Error al eliminar la planificación")
            print(f"❌ Error eliminando planificación {plan_id} para usuario {current_user['email']}")
        
        return redirect(url_for('mi_calendario'))
        
    except Exception as e:
        print(f"❌ Error en eliminar_planificacion: {e}")
        flash("Error interno al eliminar la planificación")
        return redirect(url_for('mi_calendario'))

@app.errorhandler(500)
def error_interno(e):
    return render_template('500.html'), 500

@app.route("/sitemap.xml")
def sitemap():
    """Generar sitemap.xml para SEO"""
    from datetime import datetime
    
    # URLs principales del sitio
    urls = [
        {
            'loc': url_for('index', _external=True),
            'lastmod': datetime.utcnow().strftime('%Y-%m-%d'),
            'changefreq': 'weekly',
            'priority': '1.0'
        },
        {
            'loc': url_for('login', _external=True),
            'lastmod': datetime.utcnow().strftime('%Y-%m-%d'),
            'changefreq': 'monthly',
            'priority': '0.8'
        },
        {
            'loc': url_for('registro', _external=True),
            'lastmod': datetime.utcnow().strftime('%Y-%m-%d'),
            'changefreq': 'monthly',
            'priority': '0.8'
        },
        {
            'loc': url_for('generar', _external=True),
            'lastmod': datetime.utcnow().strftime('%Y-%m-%d'),
            'changefreq': 'weekly',
            'priority': '0.9'
        },
        {
            'loc': url_for('planificacion', _external=True),
            'lastmod': datetime.utcnow().strftime('%Y-%m-%d'),
            'changefreq': 'weekly',
            'priority': '0.9'
        },
        {
            'loc': url_for('como_funciona', _external=True),
            'lastmod': datetime.utcnow().strftime('%Y-%m-%d'),
            'changefreq': 'monthly',
            'priority': '0.7'
        }
    ]
    
    sitemap_xml = render_template('sitemap.xml', urls=urls)
    response = make_response(sitemap_xml)
    response.headers["Content-Type"] = "application/xml"
    return response

@app.route("/robots.txt")
def robots():
    """Servir robots.txt para SEO"""
    return send_from_directory('static', 'robots.txt')

@app.route("/google48eb92cb7318a041.html")
def google_verification():
    """Archivo de verificación de Google Search Console"""
    return send_from_directory('static', 'google48eb92cb7318a041.html')

# =====================================================
# SISTEMA DE CARPETAS PARA ORGANIZAR EXÁMENES
# =====================================================

@app.route("/carpetas")
def carpetas():
    """Listar carpetas del usuario"""
    if not is_authenticated():
        return redirect(url_for('login'))
    
    try:
        if supabase:
            # Obtener carpetas del usuario
            carpetas_response = supabase.table('carpetas').select('*').eq('usuario_id', session.get('user_id')).order('fecha_creacion', desc=True).execute()
            
            # Obtener estadísticas de cada carpeta
            carpetas_con_stats = []
            for carpeta in carpetas_response.data:
                # Contar exámenes en esta carpeta
                examenes_response = supabase.table('examenes').select('id', count='exact').eq('carpeta_id', carpeta['id']).execute()
                cantidad_examenes = examenes_response.count if hasattr(examenes_response, 'count') else 0
                
                # Convertir fechas de string a datetime
                from datetime import datetime
                if carpeta.get('fecha_creacion'):
                    try:
                        carpeta['fecha_creacion'] = datetime.fromisoformat(carpeta['fecha_creacion'].replace('Z', '+00:00'))
                    except:
                        carpeta['fecha_creacion'] = None
                
                if carpeta.get('fecha_actualizacion'):
                    try:
                        carpeta['fecha_actualizacion'] = datetime.fromisoformat(carpeta['fecha_actualizacion'].replace('Z', '+00:00'))
                    except:
                        carpeta['fecha_actualizacion'] = None
                
                carpeta['cantidad_examenes'] = cantidad_examenes
                carpetas_con_stats.append(carpeta)
            
            return render_template("carpetas.html", carpetas=carpetas_con_stats)
            
    except Exception as e:
        print(f"❌ Error obteniendo carpetas: {e}")
        flash("Error al cargar las carpetas")
        return redirect(url_for('generar'))
    
    return render_template("carpetas.html", carpetas=[])

@app.route("/carpetas/crear", methods=["GET", "POST"])
def crear_carpeta():
    """Crear nueva carpeta"""
    if not is_authenticated():
        return redirect(url_for('login'))
    
    if request.method == "POST":
        nombre = request.form.get("nombre")
        descripcion = request.form.get("descripcion")
        color = request.form.get("color", "#10a37f")
        
        if not nombre:
            flash("El nombre de la carpeta es obligatorio")
            return render_template("crear_carpeta.html")
        
        try:
            if supabase:
                nueva_carpeta = {
                    'usuario_id': session.get('user_id'),
                    'nombre': nombre,
                    'descripcion': descripcion,
                    'color': color
                }
                
                response = supabase.table('carpetas').insert(nueva_carpeta).execute()
                
                if response.data:
                    flash("Carpeta creada exitosamente!")
                    return redirect(url_for('carpetas'))
                else:
                    flash("Error al crear la carpeta")
                    
        except Exception as e:
            print(f"❌ Error creando carpeta: {e}")
            flash("Error al crear la carpeta")
    
    return render_template("crear_carpeta.html")

@app.route("/carpetas/<carpeta_id>")
def ver_carpeta(carpeta_id):
    """Ver carpeta específica y sus exámenes"""
    if not is_authenticated():
        return redirect(url_for('login'))
    
    try:
        if supabase:
            # Obtener información de la carpeta
            carpeta_response = supabase.table('carpetas').select('*').eq('id', carpeta_id).eq('usuario_id', session.get('user_id')).execute()
            
            if not carpeta_response.data:
                flash("Carpeta no encontrada")
                return redirect(url_for('carpetas'))
            
            carpeta = carpeta_response.data[0]
            
            # Convertir fechas de string a datetime
            from datetime import datetime
            if carpeta.get('fecha_creacion'):
                try:
                    carpeta['fecha_creacion'] = datetime.fromisoformat(carpeta['fecha_creacion'].replace('Z', '+00:00'))
                except:
                    carpeta['fecha_creacion'] = None
            
            if carpeta.get('fecha_actualizacion'):
                try:
                    carpeta['fecha_actualizacion'] = datetime.fromisoformat(carpeta['fecha_actualizacion'].replace('Z', '+00:00'))
                except:
                    carpeta['fecha_actualizacion'] = None
            
            # Obtener exámenes de esta carpeta
            examenes_response = supabase.table('examenes').select('*').eq('carpeta_id', carpeta_id).order('fecha_creacion', desc=True).execute()
            
            # Convertir fechas de los exámenes también
            examenes = []
            for examen in examenes_response.data:
                if examen.get('fecha_creacion'):
                    try:
                        examen['fecha_creacion'] = datetime.fromisoformat(examen['fecha_creacion'].replace('Z', '+00:00'))
                    except:
                        examen['fecha_creacion'] = None
                examenes.append(examen)
            
            return render_template("carpeta_detalle.html", carpeta=carpeta, examenes=examenes)
            
    except Exception as e:
        print(f"❌ Error obteniendo carpeta: {e}")
        flash("Error al cargar la carpeta")
        return redirect(url_for('carpetas'))
    
    return redirect(url_for('carpetas'))

@app.route("/carpetas/<carpeta_id>/editar", methods=["GET", "POST"])
def editar_carpeta(carpeta_id):
    """Editar carpeta existente"""
    if not is_authenticated():
        return redirect(url_for('login'))
    
    try:
        if supabase:
            if request.method == "POST":
                nombre = request.form.get("nombre")
                descripcion = request.form.get("descripcion")
                color = request.form.get("color", "#10a37f")
                
                if not nombre:
                    flash("El nombre de la carpeta es obligatorio")
                    return render_template("editar_carpeta.html", carpeta=carpeta)
                
                # Actualizar carpeta
                update_data = {
                    'nombre': nombre,
                    'descripcion': descripcion,
                    'color': color
                }
                
                response = supabase.table('carpetas').update(update_data).eq('id', carpeta_id).eq('usuario_id', session.get('user_id')).execute()
                
                if response.data:
                    flash("Carpeta actualizada exitosamente!")
                    return redirect(url_for('ver_carpeta', carpeta_id=carpeta_id))
                else:
                    flash("Error al actualizar la carpeta")
            else:
                # GET: mostrar formulario de edición
                carpeta_response = supabase.table('carpetas').select('*').eq('id', carpeta_id).eq('usuario_id', session.get('user_id')).execute()
                
                if not carpeta_response.data:
                    flash("Carpeta no encontrada")
                    return redirect(url_for('carpetas'))
                
                carpeta = carpeta_response.data[0]
                return render_template("editar_carpeta.html", carpeta=carpeta)
                
    except Exception as e:
        print(f"❌ Error editando carpeta: {e}")
        flash("Error al editar la carpeta")
        return redirect(url_for('carpetas'))
    
    return redirect(url_for('carpetas'))

@app.route("/carpetas/<carpeta_id>/eliminar", methods=["POST"])
def eliminar_carpeta(carpeta_id):
    """Eliminar carpeta"""
    if not is_authenticated():
        return redirect(url_for('login'))
    
    try:
        if supabase:
            # Verificar que la carpeta pertenece al usuario
            carpeta_response = supabase.table('carpetas').select('id').eq('id', carpeta_id).eq('usuario_id', session.get('user_id')).execute()
            
            if not carpeta_response.data:
                flash("Carpeta no encontrada")
                return redirect(url_for('carpetas'))
            
            # Eliminar carpeta (los exámenes se quedarán sin carpeta por ON DELETE SET NULL)
            response = supabase.table('carpetas').delete().eq('id', carpeta_id).execute()
            
            if response.data:
                flash("Carpeta eliminada exitosamente!")
            else:
                flash("Error al eliminar la carpeta")
                
    except Exception as e:
        print(f"❌ Error eliminando carpeta: {e}")
        flash("Error al eliminar la carpeta")
    
    return redirect(url_for('carpetas'))

# =====================================================
# SISTEMA SIMPLE DE AUTENTICACIÓN CON SUPABASE AUTH
# =====================================================

def is_authenticated():
    """Verificar si el usuario está autenticado"""
    has_user_id = 'user_id' in session
    has_user_email = 'user_email' in session
    print(f"🔍 is_authenticated() - user_id: {has_user_id}, user_email: {has_user_email}")
    return has_user_id and has_user_email

def get_current_user():
    """Obtener datos del usuario actual desde la sesión"""
    if is_authenticated():
        return {
            'id': session.get('user_id'),
            'email': session.get('user_email'),
            'nombre': session.get('user_nombre')
        }
    return None

# Headers de seguridad básicos
@app.after_request
def add_security_headers(response):
    response.headers['X-Frame-Options'] = 'DENY'  # Previene clickjacking
    response.headers['X-Content-Type-Options'] = 'nosniff'  # Previene MIME sniffing
    response.headers['X-XSS-Protection'] = '1; mode=block'  # Protección XSS básica
    response.headers['Referrer-Policy'] = 'strict-origin-when-cross-origin'  # Control de referrer
    return response

if __name__ == '__main__':
    # Configuración para desarrollo
    debug = os.environ.get('FLASK_ENV') == 'development'
    port = int(os.environ.get('PORT', 8080))
    app.run(host='0.0.0.0', port=port, debug=debug)